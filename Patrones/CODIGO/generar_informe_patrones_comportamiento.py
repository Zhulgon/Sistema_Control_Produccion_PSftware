from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt


BASE_DIR = Path(__file__).resolve().parent
OUTPUT_PATH = BASE_DIR / "Informe_Patrones_Comportamiento_MES_Software.docx"
UML_DIR = BASE_DIR / "uml_comportamiento"


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)

    for style_name, size in (("Heading 1", 14), ("Heading 2", 12)):
        style = document.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True


def add_paragraph(document: Document, text: str, *, bold: bool = False, center: bool = False) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def add_bullet(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.line_spacing = 1.5
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def add_heading(document: Document, text: str, level: int = 1) -> None:
    heading = document.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)

    for row_values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_values):
            cells[idx].text = value
            for run in cells[idx].paragraphs[0].runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(10)


def add_pattern_section(
    document: Document,
    title: str,
    purpose: str,
    code_evidence: list[str],
    importance_points: list[str],
) -> None:
    add_heading(document, title, level=2)
    add_paragraph(document, purpose)
    add_paragraph(document, "Evidencia concreta dentro del proyecto:")
    for item in code_evidence:
        add_bullet(document, item)
    add_paragraph(document, "Importancia para el escalamiento del sistema:")
    for item in importance_points:
        add_bullet(document, item)


def build_document() -> None:
    document = Document()
    configure_document(document)

    for _ in range(4):
        add_paragraph(document, "", center=True)
    add_paragraph(document, "DOCUMENTO COMPLEMENTARIO DE PATRONES DE COMPORTAMIENTO", bold=True, center=True)
    add_paragraph(document, "Evidencia técnica de Observer, Strategy, Command y State en el proyecto MES", bold=True, center=True)
    for _ in range(3):
        add_paragraph(document, "", center=True)
    add_paragraph(document, "Asignatura: Patrones de Software", center=True)
    add_paragraph(document, "Programa: Ingeniería de Sistemas", center=True)
    add_paragraph(document, "Integrantes: Juan Sebastian Rubiano Romero - Angie Lucero Vargas Amado", center=True)
    add_paragraph(document, "Proyecto: Sistema de Control de Produccion (MES)", center=True)
    add_paragraph(document, f"Fecha de elaboración: {datetime.now().strftime('%Y-%m-%d')}", center=True)
    document.add_page_break()

    add_heading(document, "1. Proposito del documento")
    add_paragraph(
        document,
        "Este documento complementa los informes previos del repositorio y se enfoca en demostrar de forma clara "
        "cómo se usan los patrones de comportamiento recientemente consolidados en la fase software del proyecto "
        "MES. La intención no es repetir toda la arquitectura del sistema, sino explicar por qué Observer, "
        "Strategy, Command y State son necesarios, dónde aparecen en el código y qué problema real resuelven.",
    )

    add_heading(document, "2. Contexto del escalamiento")
    add_paragraph(
        document,
        "El proyecto ya contaba con persistencia SQLite, interfaz gráfica, autenticación básica, consultas de "
        "negocio y una base funcional de patrones creacionales y estructurales. Sin embargo, para sostener el "
        "escalamiento hacia una solución más clara, mantenible y defendible, se requería fortalecer la capa de "
        "comportamiento. En especial, era importante asegurar transiciones de orden bien definidas, reacciones "
        "desacopladas, consultas extensibles y operaciones encapsuladas.",
    )

    add_heading(document, "3. Patrones de comportamiento incorporados")
    add_table(
        document,
        ["Patrón", "Problema que resuelve", "Clases principales", "Impacto directo"],
        [
            [
                "State",
                "Evitar condicionales frágiles en el ciclo de vida de la orden.",
                "OrderLifecycleContext, OrderState, CreatedState, PlannedState, DispatchedState, RunningState, CompletedState",
                "Controla el flujo CREATED -> PLANNED -> DISPATCHED -> RUNNING -> COMPLETED.",
            ],
            [
                "Observer",
                "Separar auditoría y notificaciones del flujo principal.",
                "OrderObserver, OrderEventSubject, AuditRepositoryObserver, NotificationObserver",
                "Cada cambio de estado genera evidencia sin acoplar la lógica de la orden.",
            ],
            [
                "Strategy",
                "Escalar la generación de consultas sin duplicar lógica fija.",
                "ConsultationStrategy, PlanningConsultationStrategy, DispatchConsultationStrategy, CapacityConsultationStrategy, ExecutionConsultationStrategy, TraceabilityConsultationStrategy, ConsultationStrategyCatalog",
                "Permite mantener y extender las 5 consultas persistentes del sistema.",
            ],
            [
                "Command",
                "Uniformar la ejecución de acciones principales del software.",
                "MESCommand, ExecuteOrderCommand, GenerateConsultationsCommand, MESCommandInvoker",
                "Encapsula la ejecución de órdenes y consultas y deja trazabilidad de comandos.",
            ],
        ],
    )

    add_heading(document, "4. Evidencia por patron")
    add_pattern_section(
        document,
        "4.1 State",
        "El patrón State modela el ciclo de vida de la orden de producción. En lugar de controlar las etapas con "
        "if/else dispersos, el contexto de la orden delega la transición al estado actual. Así se conserva una "
        "regla clara de negocio: cada estado sabe a cuál estado puede avanzar y en qué momento.",
        [
            "Se implementa en software_mes_behavioral.py mediante OrderState y sus estados concretos.",
            "OrderLifecycleContext centraliza el estado actual, el historial y la notificación de eventos.",
            "Las transiciones se ejecutan con transition_to(), donde también se construye el registro histórico.",
            "En software_mes_service.py el flujo operativo ejecuta plan(), dispatch(), start() y complete().",
        ],
        [
            "Hace que la orden sea más fácil de explicar en sustentación porque el flujo queda visible y secuencial.",
            "Reduce errores al evitar transiciones inválidas desde estados incorrectos.",
            "Facilita persistir el historial de cambios y soporta trazabilidad real en la base de datos.",
        ],
    )
    add_pattern_section(
        document,
        "4.2 Observer",
        "Observer permite que el cambio de estado de una orden dispare acciones adicionales sin mezclar todo en la "
        "misma clase. El contexto solo notifica el evento; luego cada observador decide qué hacer con él.",
        [
            "OrderEventSubject administra attach() y notify() para la lista de observadores.",
            "AuditRepositoryObserver guarda evidencia en audit_events cada vez que hay una transición.",
            "NotificationObserver construye mensajes legibles para la interfaz y el seguimiento funcional.",
            "OrderLifecycleContext notifica automáticamente en transition_to() después de cambiar de estado.",
        ],
        [
            "Desacopla la lógica principal de la orden de responsabilidades secundarias como auditoría o mensajes.",
            "Permite agregar nuevos observadores después sin reescribir el flujo de ejecución principal.",
            "Fortalece la explicación del proyecto como software trazable y no solo como simulación.",
        ],
    )
    add_pattern_section(
        document,
        "4.3 Strategy",
        "Strategy se usa para construir las consultas del sistema como variantes intercambiables. Cada estrategia "
        "conoce su objetivo, su pregunta de negocio, el SQL que necesita y el tipo de evidencia que produce.",
        [
            "ConsultationStrategy define el contrato build(database, filters, plant_capacity_units).",
            "Las clases PlanningConsultationStrategy, DispatchConsultationStrategy, CapacityConsultationStrategy, ExecutionConsultationStrategy y TraceabilityConsultationStrategy implementan cada consulta.",
            "ConsultationStrategyCatalog reúne las estrategias y las entrega a la capa de servicio.",
            "ScalableMESProject usa strategy_catalog.build_all() al generar las 5 consultas persistentes.",
        ],
        [
            "Hace que agregar una sexta consulta sea una extensión controlada y no un bloque gigante de código repetido.",
            "Mantiene separada la lógica de cada consulta, lo que mejora lectura, pruebas y mantenimiento.",
            "Ayuda a explicar que las consultas no son casos aislados, sino una familia de comportamientos analíticos.",
        ],
    )
    add_pattern_section(
        document,
        "4.4 Command",
        "Command encapsula acciones importantes del sistema como objetos ejecutables. En este caso se aplica para "
        "uniformar la ejecución de órdenes y la generación de consultas, dejando además una historia de comandos "
        "ejecutados por el software.",
        [
            "MESCommand define la operación execute() como contrato común.",
            "ExecuteOrderCommand encapsula el flujo de ejecutar una orden con usuario y request.",
            "GenerateConsultationsCommand encapsula la construcción de consultas filtradas.",
            "MESCommandInvoker ejecuta el comando y conserva la historia de operaciones ejecutadas.",
            "ScalableMESProject envuelve execute_order() y generate_filtered_consultations() usando el invoker.",
        ],
        [
            "Ordena mejor la capa de servicio y deja claro que ciertas acciones del sistema son comandos de negocio.",
            "Permite agregar auditoría o políticas futuras alrededor de run(command) sin tocar cada caso por separado.",
            "Refuerza el carácter escalable del proyecto al formalizar operaciones que antes podían quedar dispersas.",
        ],
    )

    add_heading(document, "5. Flujo integrado de los cuatro patrones")
    add_paragraph(
        document,
        "Los cuatro patrones no trabajan de forma aislada. Durante la ejecución de una orden, Command encapsula la "
        "operación de alto nivel, State controla cada fase de la orden y Observer reacciona a cada transición para "
        "persistir auditoría y generar mensajes. Posteriormente, al solicitar las consultas, otro Command activa la "
        "generación y Strategy decide cómo construir cada una de las 5 consultas del sistema.",
    )
    for item in [
        "ExecuteOrderCommand invoca el flujo de negocio completo para una orden.",
        "OrderLifecycleContext aplica State para pasar por las etapas del proceso.",
        "OrderEventSubject notifica a AuditRepositoryObserver y NotificationObserver en cada cambio.",
        "GenerateConsultationsCommand activa la generación del bloque analítico.",
        "ConsultationStrategyCatalog selecciona y ejecuta las estrategias de consulta.",
        "Los resultados quedan persistidos en consultation_batches y consultation_results.",
    ]:
        add_bullet(document, item)

    add_heading(document, "6. Importancia de estos patrones dentro del proyecto")
    add_table(
        document,
        ["Criterio", "Aporte de los patrones"],
        [
            ["Mantenibilidad", "El código queda dividido por responsabilidades y no por bloques monolíticos."],
            ["Escalabilidad", "Es posible agregar nuevas transiciones, observadores, consultas o comandos con menor impacto."],
            ["Trazabilidad", "State y Observer dejan evidencia directa en historial y auditoría persistida."],
            ["Claridad académica", "Cada patrón se puede mostrar con un caso real y no solo con clases teóricas."],
            ["Sustentación", "La integración de los cuatro patrones facilita explicar decisiones de diseño con ejemplos del sistema."],
        ],
    )

    add_heading(document, "7. Evidencia de UML requerida")
    add_paragraph(
        document,
        "Sí, los UML son necesarios y altamente recomendables para este entregable. Aunque el código ya demuestra "
        "la implementación, los diagramas permiten visualizar la relación entre clases, interfaces, dependencias y "
        "puntos de extensión. Para la sustentación, los UML ayudan a explicar de forma rápida por qué cada patrón "
        "existe y cómo encaja dentro del software MES.",
    )
    add_paragraph(document, "En esta entrega se dejan preparados los siguientes archivos PlantUML:")
    for item in [
        f"State: {UML_DIR / 'state_mesoftware.puml'}",
        f"Observer: {UML_DIR / 'observer_mesoftware.puml'}",
        f"Strategy: {UML_DIR / 'strategy_mesoftware.puml'}",
        f"Command: {UML_DIR / 'command_mesoftware.puml'}",
    ]:
        add_bullet(document, item)

    add_heading(document, "8. Relacion con la interfaz y la persistencia")
    add_paragraph(
        document,
        "El valor de estos patrones aumenta porque no quedan aislados en una capa teórica. La interfaz principal "
        "UI_MES_Software_Pro.py muestra el resultado del historial de estados, las notificaciones del Observer y la "
        "recarga del historial persistido de consultas. A su vez, SQLite conserva las órdenes, los eventos y los "
        "lotes de consultas, lo que demuestra que los patrones introducidos afectan el comportamiento real del sistema.",
    )

    add_heading(document, "9. Verificacion realizada")
    for item in [
        "Se ejecutó la prueba de integración test_mes_software_upgrade.py y pasó correctamente.",
        "La prueba valida ejecución de orden, persistencia de estado, persistencia de consultas y recarga del historial.",
        "También verifica que el mapa de patrones del proyecto ya incluye Observer, Strategy, Command y State.",
        "Las pruebas conceptuales históricas del repositorio continuaron pasando después de la integración.",
    ]:
        add_bullet(document, item)

    add_heading(document, "10. Conclusiones")
    add_paragraph(
        document,
        "La incorporación de Observer, Strategy, Command y State fortalece la fase software del proyecto porque "
        "convierte comportamientos clave del MES en estructuras de diseño explicables, extensibles y mantenibles. "
        "Estos patrones no se agregaron como relleno académico; cada uno responde a una necesidad concreta del "
        "sistema: controlar estados, reaccionar a eventos, escalar consultas y encapsular operaciones.",
    )
    add_paragraph(
        document,
        "Con este complemento, el proyecto queda mejor preparado para sustentación, ya que ahora existe evidencia "
        "documental, evidencia en código, evidencia en persistencia y evidencia visual mediante UML para los cuatro "
        "patrones de comportamiento trabajados en esta etapa.",
    )

    document.save(str(OUTPUT_PATH))
    print(f"Documento generado en: {OUTPUT_PATH}")


if __name__ == "__main__":
    build_document()
