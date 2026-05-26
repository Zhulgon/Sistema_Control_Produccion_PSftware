from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt


OUTPUT_PATH = (
    Path(__file__).resolve().parents[1]
    / "Documentos 3 corte"
    / "DOCUMENTOS"
    / "Consultas"
    / "Informe_Consultas_Adicionales_MES_Software.docx"
)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)

    for style_name, size in (("Heading 1", 14), ("Heading 2", 12)):
        style = document.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True


def add_paragraph(document: Document, text: str, *, bold: bool = False, center: bool = False) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def add_bullet(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.line_spacing = 1.5
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def add_heading(document: Document, text: str, level: int = 1) -> None:
    heading = document.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)

    for row_values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_values):
            cells[idx].text = value
            for run in cells[idx].paragraphs[0].runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(10)


def add_consultation_section(
    document: Document,
    title: str,
    objective: str,
    question: str,
    filters: list[str],
    tables: list[str],
    value: str,
) -> None:
    add_heading(document, title, level=2)
    add_paragraph(document, f"Objetivo: {objective}")
    add_paragraph(document, f"Pregunta de negocio: {question}")
    add_paragraph(document, "Filtros sugeridos:")
    for item in filters:
        add_bullet(document, item)
    add_paragraph(document, "Fuentes de datos involucradas:")
    for item in tables:
        add_bullet(document, item)
    add_paragraph(document, f"Valor dentro del proyecto: {value}")


def build_document() -> None:
    document = Document()
    configure_document(document)

    for _ in range(4):
        add_paragraph(document, "", center=True)
    add_paragraph(document, "PROPUESTA E IMPLEMENTACION DE CONSULTAS ADICIONALES", bold=True, center=True)
    add_paragraph(document, "Ampliacion del bloque analitico del sistema MES de 5 a 10 consultas", bold=True, center=True)
    for _ in range(3):
        add_paragraph(document, "", center=True)
    add_paragraph(document, "Asignatura: Patrones de Software", center=True)
    add_paragraph(document, "Programa: Ingeniería de Sistemas", center=True)
    add_paragraph(document, "Integrantes: Juan Sebastian Rubiano Romero - Angie Lucero Vargas Amado", center=True)
    add_paragraph(document, "Proyecto: Sistema de Control de Producción (MES)", center=True)
    add_paragraph(document, f"Fecha de elaboración: {datetime.now().strftime('%Y-%m-%d')}", center=True)
    document.add_page_break()
    add_heading(document, "1. Proposito del documento")
    add_paragraph(
        document,
        "Este documento presenta la consolidacion del bloque analitico del sistema MES en un conjunto de 10 consultas de negocio. La intencion ya no es solo ampliar la cantidad de consultas, sino asegurar que cada una aporte informacion util para sustentar el valor operativo y comercial de la solucion frente a un posible comprador o adoptante.",
    )
    add_heading(document, "2. Justificacion de la ampliacion")
    add_paragraph(
        document,
        "La ampliacion a 10 consultas se reviso bajo un criterio de utilidad real. Se conservaron las consultas que ayudan a explicar demanda, capacidad, cumplimiento, trazabilidad, productividad, seguridad y desempeno operativo. Ademas, se sustituyo la lectura menos convincente para negocio por una consulta de desempeno por producto, ya que esta aporta mejor soporte para decisiones de portafolio, priorizacion comercial y analisis de rentabilidad operativa.",
    )
    add_heading(document, "3. Resumen de las nuevas consultas")
    add_table(
        document,
        ["ID", "Nombre", "Enfoque principal", "Utilidad"],
        [
            ["SQ6", "Productividad por operador", "Rendimiento humano", "Comparar desempeno y oportunidades de mejora por operador."],
            ["SQ7", "Carga y utilizacion por linea", "Balance de recursos", "Detectar lineas con mayor concentracion de trabajo."],
            ["SQ8", "Seguridad operativa y auditoria", "Gobierno operativo", "Rastrear accesos y acciones relevantes del sistema."],
            ["SQ9", "Desempeno por producto", "Portafolio operativo", "Priorizar productos por volumen, cumplimiento y eficiencia."],
            ["SQ10", "Comparativo de desempeno por turno", "Comparacion operativa", "Detectar diferencias entre DAY y NIGHT para intervenir mejor la operacion."],
        ],
    )
    add_heading(document, "4. Desarrollo de las consultas")
    add_consultation_section(
        document,
        "4.1 SQ6 - Consulta de productividad por operador",
        "Comparar el aporte operativo por operador segun ordenes cerradas, volumen producido, tiempos de parada y eficiencia promedio.",
        "Que operadores entregan mejor productividad y donde hay oportunidad de mejora?",
        [
            "Rango de fecha y hora",
            "Linea de produccion",
            "Operator ID",
            "Turno",
            "Protocolo",
            "Estado",
        ],
        [
            "Tabla orders",
        ],
        "Esta consulta ayuda a justificar que el MES no solo controla maquinas, sino que tambien facilita gestion del desempeno operativo del equipo humano.",
    )
    add_consultation_section(
        document,
        "4.2 SQ7 - Consulta de carga y utilizacion por linea",
        "Mostrar como se distribuye la demanda entre lineas para balancear carga y priorizar recursos.",
        "Que lineas concentran mas trabajo y como se esta utilizando la capacidad instalada?",
        [
            "Rango de fecha y hora",
            "Linea de produccion",
            "Producto",
            "Turno",
            "Estado",
        ],
        [
            "Tabla orders",
        ],
        "Aporta una lectura clara sobre saturacion de lineas y capacidad de respuesta, dos argumentos importantes para sustentar valor operativo.",
    )
    add_consultation_section(
        document,
        "4.3 SQ8 - Consulta de seguridad operativa y auditoria",
        "Rastrear eventos de login, transiciones de estado y acciones auditadas sobre ordenes, usuarios y lineas.",
        "Que usuarios, ordenes o lineas presentan actividad auditada relevante en el periodo?",
        [
            "Rango de fecha y hora",
            "Order ID",
            "Linea de produccion",
            "Usuario solicitante",
            "Estado",
        ],
        [
            "Tabla audit_events",
            "Tabla orders",
        ],
        "Fortalece la confianza en la solucion porque demuestra control, evidencia y gobierno basico de la operacion.",
    )
    add_consultation_section(
        document,
        "4.4 SQ9 - Consulta de desempeno por producto",
        "Comparar volumen, cumplimiento, desviacion, paradas y eficiencia por producto para priorizar el portafolio operativo.",
        "Que productos mueven mas volumen y cuales presentan peor eficiencia o mayor desviacion?",
        [
            "Rango de fecha y hora",
            "Producto",
            "Linea de produccion",
            "Turno",
            "Estado",
        ],
        [
            "Tabla orders",
        ],
        "Esta consulta reemplaza una lectura menos comercial y mejora la capacidad del sistema para apoyar decisiones sobre portafolio, foco operativo y priorizacion de mejora.",
    )
    add_consultation_section(
        document,
        "4.5 SQ10 - Consulta comparativa por turno",
        "Comparar resultados entre los turnos DAY y NIGHT mediante produccion, paradas y eficiencia promedio.",
        "Que diferencias operativas hay entre DAY y NIGHT y donde conviene intervenir?",
        [
            "Rango de fecha y hora",
            "Turno",
            "Linea de produccion",
            "Operator ID",
            "Estado",
        ],
        [
            "Tabla orders",
        ],
        "Permite demostrar que el sistema soporta comparaciones utiles para programacion, supervision y mejora continua.",
    )
    add_heading(document, "5. Relacion con la implementacion")
    add_paragraph(
        document,
        "Estas cinco consultas quedaron integradas al proyecto con el mismo enfoque arquitectonico usado en las anteriores. Cada una se construye como una estrategia independiente dentro del patron Strategy y luego es orquestada por la capa de servicio para persistirse en la base SQLite junto con sus filtros, parametros, SQL sugerido y muestra de resultados.",
    )
    for item in [
        "SQ6 se implementa como OperatorProductivityConsultationStrategy.",
        "SQ7 se implementa como LineLoadConsultationStrategy.",
        "SQ8 se implementa como AuditActivityConsultationStrategy.",
        "SQ9 se implementa como ProductPerformanceConsultationStrategy.",
        "SQ10 se implementa como ShiftPerformanceConsultationStrategy.",
    ]:
        add_bullet(document, item)
    add_heading(document, "6. Conclusion")
    add_paragraph(
        document,
        "La ampliacion del sistema a 10 consultas mejora la utilidad analitica del proyecto y lo hace mas convincente como producto de software. El bloque resultante combina operacion, seguridad, trazabilidad, integracion y portafolio de productos, lo que permite presentar una solucion MES con argumentos mas claros para gestion, supervision y adopcion.",
    )
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(OUTPUT_PATH))
    print(f"Documento generado en: {OUTPUT_PATH}")


if __name__ == "__main__":
    build_document()
