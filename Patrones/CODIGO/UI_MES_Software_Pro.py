"""
Interfaz profesional para la evolucion software del proyecto MES.

Ejecucion:
python UI_MES_Software_Pro.py
"""

from __future__ import annotations

import json
import sys
import tkinter as tk
from datetime import datetime
from pathlib import Path
from tkinter import filedialog, messagebox, ttk
from tkinter.scrolledtext import ScrolledText

from docx import Document

CURRENT_DIR = str(Path(__file__).resolve().parent)
if CURRENT_DIR not in sys.path:
    sys.path.insert(0, CURRENT_DIR)

from mes_functional_app import OrderExecutionRequest
from software_mes_queries import QueryFilters
from software_mes_service import ScalableMESProject


class VerticalScrolledPanel(ttk.Frame):
    def __init__(
        self,
        parent: ttk.Widget,
        *,
        style: str = "App.TFrame",
        padding: int | tuple[int, ...] = 0,
        canvas_background: str = "#edf2f7",
    ) -> None:
        super().__init__(parent, style=style)
        self.columnconfigure(0, weight=1)
        self.rowconfigure(0, weight=1)

        self._canvas = tk.Canvas(
            self,
            background=canvas_background,
            highlightthickness=0,
            borderwidth=0,
        )
        self._scrollbar = ttk.Scrollbar(self, orient=tk.VERTICAL, command=self._canvas.yview)
        self._canvas.configure(yscrollcommand=self._scrollbar.set)

        self._canvas.grid(row=0, column=0, sticky="nsew")
        self._scrollbar.grid(row=0, column=1, sticky="ns", padx=(8, 0))

        self.body = ttk.Frame(self._canvas, style=style, padding=padding)
        self._body_window = self._canvas.create_window((0, 0), window=self.body, anchor="nw")

        self.body.bind("<Configure>", self._sync_scroll_region)
        self._canvas.bind("<Configure>", self._resize_body)

    def reset_view(self) -> None:
        self._canvas.yview_moveto(0)

    def _sync_scroll_region(self, _event: tk.Event) -> None:
        self._canvas.configure(scrollregion=self._canvas.bbox("all"))

    def _resize_body(self, event: tk.Event) -> None:
        self._canvas.itemconfigure(self._body_window, width=event.width)


class MESSoftwareControlCenter(tk.Tk):
    def __init__(self) -> None:
        super().__init__()
        self.title("MES Control Center | Software Edition")
        self.geometry("1440x860")
        self.minsize(1180, 680)
        self.configure(background="#edf2f7")

        self.service = ScalableMESProject()
        self.current_user = None
        self.current_result: dict | None = None
        self.current_consultations: list[dict] = []
        self.pattern_map: list[dict] = self.service.pattern_map()
        self.status_var = tk.StringVar(value="Inicia sesion para habilitar la operacion del sistema.")
        self.session_var = tk.StringVar(value="Sesion no iniciada")
        self.db_var = tk.StringVar(value=self.service.db_path)

        self._init_vars()
        self._build_theme()
        self._build_layout()
        self._bind_events()
        self._load_demo_defaults()
        self._lock_workspace()

    def _init_vars(self) -> None:
        self.username_var = tk.StringVar()
        self.password_var = tk.StringVar()
        self.template_var = tk.StringVar()
        self.template_key_var = tk.StringVar()
        self.template_name_var = tk.StringVar()
        self.machine_type_var = tk.StringVar(value="cnc")
        self.line_var = tk.StringVar()
        self.program_var = tk.StringVar()
        self.product_code_var = tk.StringVar()
        self.product_name_var = tk.StringVar()
        self.line_mode_var = tk.StringVar(value="automatic")
        self.quality_checks_var = tk.StringVar()
        self.order_id_var = tk.StringVar()
        self.units_var = tk.StringVar(value="300")
        self.shift_var = tk.StringVar(value="DAY")
        self.protocol_var = tk.StringVar(value="opcua")
        self.operator_id_var = tk.StringVar()
        self.operator_name_var = tk.StringVar()

        today = datetime.now().date().isoformat()
        self.filter_order_id_var = tk.StringVar()
        self.filter_product_code_var = tk.StringVar()
        self.filter_line_var = tk.StringVar()
        self.filter_shift_var = tk.StringVar()
        self.filter_protocol_var = tk.StringVar()
        self.filter_operator_id_var = tk.StringVar()
        self.filter_state_var = tk.StringVar()
        self.filter_requested_by_var = tk.StringVar()
        self.filter_date_from_var = tk.StringVar(value=today)
        self.filter_date_to_var = tk.StringVar(value=today)
        self.filter_hour_from_var = tk.StringVar(value="00:00:00")
        self.filter_hour_to_var = tk.StringVar(value="23:59:59")
        self.filter_limit_var = tk.StringVar(value="25")

    def _build_theme(self) -> None:
        style = ttk.Style(self)
        style.theme_use("clam")

        style.configure(".", font=("Bahnschrift", 10), foreground="#17202a")
        style.configure("Sidebar.TFrame", background="#10233f")
        style.configure("SidebarCard.TFrame", background="#16345f")
        style.configure("App.TFrame", background="#edf2f7")
        style.configure("Card.TFrame", background="#ffffff", relief="flat")
        style.configure("SoftCard.TFrame", background="#f7fafc", relief="flat")
        style.configure("Hero.TFrame", background="#ffffff", relief="flat")
        style.configure("SidebarTitle.TLabel", background="#10233f", foreground="#f8fafc", font=("Bahnschrift", 20, "bold"))
        style.configure("SidebarText.TLabel", background="#10233f", foreground="#d7e3f4", font=("Bahnschrift", 10))
        style.configure("SidebarMuted.TLabel", background="#16345f", foreground="#d7e3f4", font=("Bahnschrift", 10))
        style.configure("HeroTitle.TLabel", background="#ffffff", foreground="#0f172a", font=("Bahnschrift", 22, "bold"))
        style.configure("HeroText.TLabel", background="#ffffff", foreground="#516173", font=("Bahnschrift", 10))
        style.configure("Section.TLabelframe", background="#ffffff", foreground="#0f172a")
        style.configure("Section.TLabelframe.Label", background="#ffffff", foreground="#0f172a", font=("Bahnschrift", 10, "bold"))
        style.configure("Accent.TButton", background="#0f766e", foreground="#ffffff", padding=(12, 8), font=("Bahnschrift", 10, "bold"))
        style.map("Accent.TButton", background=[("active", "#115e59")])
        style.configure("Ghost.TButton", background="#e2e8f0", foreground="#10233f", padding=(10, 8), font=("Bahnschrift", 10, "bold"))
        style.map("Ghost.TButton", background=[("active", "#cbd5e1")])
        style.configure("Warn.TButton", background="#b45309", foreground="#ffffff", padding=(10, 8), font=("Bahnschrift", 10, "bold"))
        style.map("Warn.TButton", background=[("active", "#92400e")])
        style.configure("MetricTitle.TLabel", background="#ffffff", foreground="#64748b", font=("Bahnschrift", 10))
        style.configure("MetricValue.TLabel", background="#ffffff", foreground="#0f172a", font=("Bahnschrift", 20, "bold"))
        style.configure("MetricHint.TLabel", background="#ffffff", foreground="#0f766e", font=("Bahnschrift", 9))
        style.configure("CardTitle.TLabel", background="#ffffff", foreground="#0f172a", font=("Bahnschrift", 12, "bold"))
        style.configure("CardHint.TLabel", background="#ffffff", foreground="#64748b", font=("Bahnschrift", 9))
        style.configure("Notebook.TNotebook", background="#edf2f7", borderwidth=0)
        style.configure("Notebook.TNotebook.Tab", padding=(14, 8), font=("Bahnschrift", 10, "bold"))
        style.map("Notebook.TNotebook.Tab", background=[("selected", "#ffffff"), ("active", "#dde6ef")])
        style.configure("Data.Treeview", rowheight=28, background="#ffffff", fieldbackground="#ffffff", foreground="#0f172a")
        style.configure("Data.Treeview.Heading", background="#dce7f3", foreground="#10233f", font=("Bahnschrift", 10, "bold"))
        style.map("Data.Treeview", background=[("selected", "#dbeafe")], foreground=[("selected", "#10233f")])

    def _build_layout(self) -> None:
        self.columnconfigure(1, weight=1)
        self.rowconfigure(0, weight=1)

        self.sidebar = ttk.Frame(self, style="Sidebar.TFrame", padding=22)
        self.sidebar.grid(row=0, column=0, sticky="nsew")
        self.sidebar.configure(width=290)
        self.sidebar.grid_propagate(False)
        self._build_sidebar()

        self.content = ttk.Frame(self, style="App.TFrame", padding=18)
        self.content.grid(row=0, column=1, sticky="nsew")
        self.content.columnconfigure(0, weight=1)
        self.content.rowconfigure(1, weight=1)

        hero = ttk.Frame(self.content, style="Hero.TFrame", padding=20)
        hero.grid(row=0, column=0, sticky="ew")
        hero.columnconfigure(0, weight=1)
        ttk.Label(hero, text="MES Software Control Center", style="HeroTitle.TLabel").grid(
            row=0, column=0, sticky="w"
        )
        ttk.Label(
            hero,
            text="Operacion de ordenes, seguridad minima, persistencia y consultas filtradas desde una sola interfaz.",
            style="HeroText.TLabel",
        ).grid(row=1, column=0, sticky="w", pady=(4, 8))
        ttk.Label(hero, textvariable=self.status_var, style="HeroText.TLabel").grid(row=2, column=0, sticky="w")
        ttk.Label(hero, textvariable=self.session_var, style="HeroText.TLabel").grid(row=0, column=1, sticky="e")

        self.notebook = ttk.Notebook(self.content, style="Notebook.TNotebook")
        self.notebook.grid(row=1, column=0, sticky="nsew", pady=(14, 0))

        self.dashboard_tab = ttk.Frame(self.notebook, style="App.TFrame", padding=6)
        self.order_tab = ttk.Frame(self.notebook, style="App.TFrame", padding=6)
        self.consultation_tab = ttk.Frame(self.notebook, style="App.TFrame", padding=6)
        self.pattern_tab = ttk.Frame(self.notebook, style="App.TFrame", padding=6)

        self.notebook.add(self.dashboard_tab, text="Dashboard")
        self.notebook.add(self.order_tab, text="Orden operativa")
        self.notebook.add(self.consultation_tab, text="Consultas")
        self.notebook.add(self.pattern_tab, text="Patrones")

        self._build_dashboard_tab()
        self._build_order_tab()
        self._build_consultation_tab()
        self._build_pattern_tab()

    def _build_sidebar(self) -> None:
        self.sidebar.columnconfigure(0, weight=1)

        ttk.Label(self.sidebar, text="MES", style="SidebarTitle.TLabel").grid(row=0, column=0, sticky="w")
        ttk.Label(self.sidebar, text="Software Edition", style="SidebarText.TLabel").grid(
            row=1, column=0, sticky="w", pady=(0, 16)
        )

        login_card = ttk.Frame(self.sidebar, style="SidebarCard.TFrame", padding=16)
        login_card.grid(row=2, column=0, sticky="ew")
        login_card.columnconfigure(0, weight=1)
        ttk.Label(login_card, text="Acceso seguro", style="SidebarMuted.TLabel").grid(row=0, column=0, sticky="w")
        ttk.Label(login_card, text="Usuario", style="SidebarMuted.TLabel").grid(row=1, column=0, sticky="w", pady=(12, 4))
        ttk.Entry(login_card, textvariable=self.username_var).grid(row=2, column=0, sticky="ew")
        ttk.Label(login_card, text="Contrasena", style="SidebarMuted.TLabel").grid(row=3, column=0, sticky="w", pady=(12, 4))
        ttk.Entry(login_card, textvariable=self.password_var, show="*").grid(row=4, column=0, sticky="ew")
        ttk.Button(login_card, text="Iniciar sesion", style="Accent.TButton", command=self._login).grid(
            row=5, column=0, sticky="ew", pady=(14, 8)
        )
        ttk.Button(login_card, text="Cerrar sesion", style="Ghost.TButton", command=self._logout).grid(
            row=6, column=0, sticky="ew"
        )

        demo_card = ttk.Frame(self.sidebar, style="SidebarCard.TFrame", padding=16)
        demo_card.grid(row=3, column=0, sticky="ew", pady=(16, 0))
        demo_card.columnconfigure(0, weight=1)
        ttk.Label(demo_card, text="Credenciales demo", style="SidebarMuted.TLabel").grid(row=0, column=0, sticky="w")
        for idx, item in enumerate(self.service.demo_credentials(), start=1):
            text = f"{item['role']} | {item['username']}"
            ttk.Button(
                demo_card,
                text=text,
                style="Ghost.TButton",
                command=lambda data=item: self._apply_demo_credentials(data),
            ).grid(row=idx, column=0, sticky="ew", pady=(8 if idx == 1 else 6, 0))

        meta_card = ttk.Frame(self.sidebar, style="SidebarCard.TFrame", padding=16)
        meta_card.grid(row=4, column=0, sticky="ew", pady=(16, 0))
        meta_card.columnconfigure(0, weight=1)
        ttk.Label(meta_card, text="Entorno", style="SidebarMuted.TLabel").grid(row=0, column=0, sticky="w")
        ttk.Label(meta_card, text="Base de datos activa", style="SidebarMuted.TLabel").grid(
            row=1, column=0, sticky="w", pady=(12, 2)
        )
        ttk.Label(meta_card, textvariable=self.db_var, style="SidebarMuted.TLabel", wraplength=240).grid(
            row=2, column=0, sticky="w"
        )
        ttk.Button(meta_card, text="Actualizar panel", style="Ghost.TButton", command=self._refresh_all_views).grid(
            row=3, column=0, sticky="ew", pady=(14, 0)
        )

    def _build_dashboard_tab(self) -> None:
        self.dashboard_tab.columnconfigure(0, weight=1)
        self.dashboard_tab.rowconfigure(1, weight=1)

        metrics = ttk.Frame(self.dashboard_tab, style="App.TFrame")
        metrics.grid(row=0, column=0, sticky="ew")
        for idx in range(6):
            metrics.columnconfigure(idx, weight=1)

        self.metric_vars = {
            "total_orders": tk.StringVar(value="0"),
            "planned_units_total": tk.StringVar(value="0"),
            "produced_units_total": tk.StringVar(value="0"),
            "downtime_minutes_total": tk.StringVar(value="0"),
            "average_oee": tk.StringVar(value="0.0%"),
            "total_consultation_batches": tk.StringVar(value="0"),
        }
        cards = [
            ("Ordenes registradas", "total_orders", "Persistidas en SQLite"),
            ("Unidades planeadas", "planned_units_total", "Capacidad comprometida"),
            ("Unidades producidas", "produced_units_total", "Resultado acumulado"),
            ("Paradas", "downtime_minutes_total", "Minutos registrados"),
            ("OEE promedio", "average_oee", "Controlado por Proxy"),
            ("Lotes de consultas", "total_consultation_batches", "Historial analitico guardado"),
        ]
        for idx, (title, key, hint) in enumerate(cards):
            card = ttk.Frame(metrics, style="Card.TFrame", padding=16)
            card.grid(row=0, column=idx, sticky="nsew", padx=(0 if idx == 0 else 8, 0))
            ttk.Label(card, text=title, style="MetricTitle.TLabel").grid(row=0, column=0, sticky="w")
            ttk.Label(card, textvariable=self.metric_vars[key], style="MetricValue.TLabel").grid(
                row=1, column=0, sticky="w", pady=(6, 4)
            )
            ttk.Label(card, text=hint, style="MetricHint.TLabel").grid(row=2, column=0, sticky="w")

        body = ttk.Frame(self.dashboard_tab, style="App.TFrame")
        body.grid(row=1, column=0, sticky="nsew", pady=(12, 0))
        body.columnconfigure(0, weight=3)
        body.columnconfigure(1, weight=2)
        body.rowconfigure(0, weight=1)

        orders_card = ttk.Frame(body, style="Card.TFrame", padding=16)
        orders_card.grid(row=0, column=0, sticky="nsew", padx=(0, 8))
        orders_card.columnconfigure(0, weight=1)
        orders_card.rowconfigure(1, weight=1)
        ttk.Label(orders_card, text="Ordenes recientes", style="CardTitle.TLabel").grid(row=0, column=0, sticky="w")
        self.recent_orders_tree = ttk.Treeview(
            orders_card,
            columns=("order", "line", "shift", "shift_window", "created_at", "state", "produced"),
            show="headings",
            style="Data.Treeview",
        )
        for key, title, width in (
            ("order", "Orden", 180),
            ("line", "Linea", 100),
            ("shift", "Turno", 80),
            ("shift_window", "Ventana turno", 180),
            ("created_at", "Registrada", 160),
            ("state", "Estado", 110),
            ("produced", "Producido", 110),
        ):
            self.recent_orders_tree.heading(key, text=title)
            self.recent_orders_tree.column(key, width=width, anchor="center")
        recent_orders_scroll = ttk.Scrollbar(orders_card, orient=tk.VERTICAL, command=self.recent_orders_tree.yview)
        self.recent_orders_tree.configure(yscrollcommand=recent_orders_scroll.set)
        self.recent_orders_tree.grid(row=1, column=0, sticky="nsew", pady=(12, 0))
        recent_orders_scroll.grid(row=1, column=1, sticky="ns", pady=(12, 0))

        audit_card = ttk.Frame(body, style="Card.TFrame", padding=16)
        audit_card.grid(row=0, column=1, sticky="nsew")
        audit_card.columnconfigure(0, weight=1)
        audit_card.rowconfigure(1, weight=1)
        ttk.Label(audit_card, text="Auditoria reciente", style="CardTitle.TLabel").grid(row=0, column=0, sticky="w")
        self.audit_text = ScrolledText(audit_card, wrap=tk.WORD, font=("Consolas", 10), relief=tk.FLAT)
        self.audit_text.grid(row=1, column=0, sticky="nsew", pady=(12, 0))
        self.audit_text.configure(state=tk.DISABLED, background="#ffffff", foreground="#0f172a")

    def _build_order_tab(self) -> None:
        self.order_tab.columnconfigure(0, weight=7)
        self.order_tab.columnconfigure(1, weight=6)
        self.order_tab.rowconfigure(0, weight=1)

        self.order_form_scroll = VerticalScrolledPanel(self.order_tab, style="App.TFrame")
        self.order_form_scroll.grid(row=0, column=0, sticky="nsew", padx=(0, 8))
        form_panel = self.order_form_scroll.body
        form_panel.columnconfigure(0, weight=1)

        template_card = ttk.LabelFrame(form_panel, text="Plantillas y producto", style="Section.TLabelframe", padding=16)
        template_card.grid(row=0, column=0, sticky="ew")
        for idx in range(4):
            template_card.columnconfigure(idx, weight=1)

        ttk.Label(template_card, text="Plantilla guardada").grid(row=0, column=0, sticky="w")
        self.template_combo = ttk.Combobox(template_card, textvariable=self.template_var, state="readonly")
        self.template_combo.grid(row=1, column=0, sticky="ew", padx=(0, 8), pady=(4, 10))
        ttk.Button(template_card, text="Cargar", style="Ghost.TButton", command=self._load_selected_template).grid(
            row=1, column=1, sticky="ew", pady=(4, 10)
        )
        ttk.Button(template_card, text="Nueva orden", style="Ghost.TButton", command=self._reset_order_form).grid(
            row=1, column=2, sticky="ew", padx=(8, 0), pady=(4, 10)
        )

        self._labeled_entry(template_card, "Codigo plantilla", self.template_key_var, 2, 0)
        self._labeled_entry(template_card, "Nombre plantilla", self.template_name_var, 2, 1)
        self._labeled_combo(template_card, "Tipo maquina", self.machine_type_var, ("cnc", "robot"), 2, 2)
        self._labeled_entry(template_card, "Linea", self.line_var, 2, 3)
        self._labeled_entry(template_card, "Programa", self.program_var, 4, 0)
        self._labeled_entry(template_card, "Codigo producto", self.product_code_var, 4, 1)
        self._labeled_entry(template_card, "Nombre producto", self.product_name_var, 4, 2)
        self._labeled_combo(template_card, "Modo linea", self.line_mode_var, ("automatic", "manual"), 4, 3)
        self._labeled_entry(template_card, "Controles calidad", self.quality_checks_var, 6, 0, colspan=2)

        ttk.Label(template_card, text="Parametros JSON").grid(row=6, column=2, sticky="w", padx=(8, 0))
        self.parameters_text = ScrolledText(template_card, wrap=tk.WORD, font=("Consolas", 10), height=6, relief=tk.FLAT)
        self.parameters_text.grid(row=7, column=2, columnspan=2, sticky="nsew", padx=(8, 0), pady=(4, 0))
        self.parameters_text.configure(background="#f8fafc", foreground="#0f172a")
        template_card.rowconfigure(7, weight=1)

        action_row = ttk.Frame(form_panel, style="App.TFrame")
        action_row.grid(row=1, column=0, sticky="ew", pady=(10, 10))
        ttk.Button(action_row, text="Guardar plantilla", style="Ghost.TButton", command=self._save_template).pack(
            side=tk.LEFT
        )

        execution_card = ttk.LabelFrame(form_panel, text="Ejecucion operativa", style="Section.TLabelframe", padding=16)
        execution_card.grid(row=2, column=0, sticky="nsew")
        for idx in range(4):
            execution_card.columnconfigure(idx, weight=1)

        self._labeled_entry(execution_card, "Order ID", self.order_id_var, 0, 0)
        self._labeled_entry(execution_card, "Unidades planeadas", self.units_var, 0, 1)
        self._labeled_combo(execution_card, "Turno", self.shift_var, ("DAY", "NIGHT"), 0, 2)
        self._labeled_combo(execution_card, "Protocolo", self.protocol_var, ("opcua", "modbus"), 0, 3)
        self._labeled_entry(execution_card, "Operator ID", self.operator_id_var, 2, 0)
        self._labeled_entry(execution_card, "Operator name", self.operator_name_var, 2, 1, colspan=2)

        execute_actions = ttk.Frame(execution_card, style="Card.TFrame")
        execute_actions.grid(row=4, column=0, columnspan=4, sticky="ew", pady=(16, 0))
        ttk.Button(execute_actions, text="Ejecutar orden", style="Accent.TButton", command=self._execute_order).pack(
            side=tk.LEFT
        )
        ttk.Button(
            execute_actions,
            text="Cargar datos demo",
            style="Warn.TButton",
            command=self._seed_demo_order,
        ).pack(side=tk.LEFT, padx=(8, 0))

        ttk.Label(
            form_panel,
            text="Usa la barra lateral de este panel si tu pantalla no alcanza a mostrar el formulario completo.",
            style="CardHint.TLabel",
            wraplength=520,
        ).grid(row=3, column=0, sticky="w", pady=(10, 0))

        output_panel = ttk.Frame(self.order_tab, style="App.TFrame")
        output_panel.grid(row=0, column=1, sticky="nsew")
        output_panel.columnconfigure(0, weight=1)
        output_panel.rowconfigure(0, weight=1)

        results_card = ttk.Frame(output_panel, style="Card.TFrame", padding=16)
        results_card.grid(row=0, column=0, sticky="nsew")
        results_card.columnconfigure(0, weight=1)
        results_card.rowconfigure(1, weight=1)

        ttk.Label(results_card, text="Resultados de la operacion", style="CardTitle.TLabel").grid(
            row=0, column=0, sticky="w"
        )
        self.result_notebook = ttk.Notebook(results_card, style="Notebook.TNotebook")
        self.result_notebook.grid(row=1, column=0, sticky="nsew", pady=(12, 0))

        self.summary_text = self._create_readonly_text(self.result_notebook)
        self.detail_text = self._create_readonly_text(self.result_notebook)
        self.history_text = self._create_readonly_text(self.result_notebook)
        self.security_text = self._create_readonly_text(self.result_notebook)

        self.result_notebook.add(self.summary_text, text="Resumen")
        self.result_notebook.add(self.detail_text, text="Operacion")
        self.result_notebook.add(self.history_text, text="Estado y eventos")
        self.result_notebook.add(self.security_text, text="Seguridad")

    def _build_consultation_tab(self) -> None:
        self.consultation_tab.columnconfigure(0, weight=1)
        self.consultation_tab.rowconfigure(1, weight=1)

        filter_card = ttk.LabelFrame(self.consultation_tab, text="Filtros de consulta", style="Section.TLabelframe", padding=16)
        filter_card.grid(row=0, column=0, sticky="ew")
        for idx in range(6):
            filter_card.columnconfigure(idx, weight=1)

        ttk.Label(
            filter_card,
            text="Bloque analitico activo: el sistema genera las 10 consultas SQ1-SQ10 definidas para la sustentacion.",
            style="CardHint.TLabel",
            wraplength=1100,
        ).grid(row=0, column=0, columnspan=6, sticky="w", pady=(0, 10))

        self._labeled_entry(filter_card, "Order ID", self.filter_order_id_var, 1, 0)
        self._labeled_entry(filter_card, "Producto", self.filter_product_code_var, 1, 1)
        self._labeled_entry(filter_card, "Linea", self.filter_line_var, 1, 2)
        self._labeled_combo(filter_card, "Turno", self.filter_shift_var, ("", "DAY", "NIGHT"), 1, 3)
        self._labeled_combo(filter_card, "Protocolo", self.filter_protocol_var, ("", "opcua", "modbus"), 1, 4)
        self._labeled_entry(filter_card, "Operator ID", self.filter_operator_id_var, 1, 5)

        self._labeled_combo(
            filter_card,
            "Estado",
            self.filter_state_var,
            ("", "CREATED", "PLANNED", "DISPATCHED", "RUNNING", "COMPLETED"),
            3,
            0,
        )
        self._labeled_entry(filter_card, "Solicitado por", self.filter_requested_by_var, 3, 1)
        self._labeled_entry(filter_card, "Fecha desde", self.filter_date_from_var, 3, 2)
        self._labeled_entry(filter_card, "Fecha hasta", self.filter_date_to_var, 3, 3)
        self._labeled_entry(filter_card, "Hora desde", self.filter_hour_from_var, 3, 4)
        self._labeled_entry(filter_card, "Hora hasta", self.filter_hour_to_var, 3, 5)
        self._labeled_entry(filter_card, "Limite", self.filter_limit_var, 5, 0)

        action_row = ttk.Frame(filter_card, style="Card.TFrame")
        action_row.grid(row=7, column=1, columnspan=5, sticky="e", pady=(16, 0))
        ttk.Button(action_row, text="Limpiar filtros", style="Ghost.TButton", command=self._reset_filters).pack(
            side=tk.RIGHT
        )
        ttk.Button(action_row, text="Exportar documento", style="Warn.TButton", command=self._export_consultations_doc).pack(
            side=tk.RIGHT, padx=(0, 8)
        )
        ttk.Button(action_row, text="Cargar historial", style="Ghost.TButton", command=self._load_persisted_consultation_history).pack(
            side=tk.RIGHT, padx=(0, 8)
        )
        ttk.Button(action_row, text="Generar consultas", style="Accent.TButton", command=self._generate_consultations).pack(
            side=tk.RIGHT, padx=(0, 8)
        )

        body = ttk.Frame(self.consultation_tab, style="App.TFrame")
        body.grid(row=1, column=0, sticky="nsew", pady=(12, 0))
        body.columnconfigure(0, weight=3)
        body.columnconfigure(1, weight=4)
        body.rowconfigure(0, weight=1)

        list_card = ttk.Frame(body, style="Card.TFrame", padding=16)
        list_card.grid(row=0, column=0, sticky="nsew", padx=(0, 8))
        list_card.columnconfigure(0, weight=1)
        list_card.rowconfigure(1, weight=1)
        ttk.Label(list_card, text="Lote de consultas", style="CardTitle.TLabel").grid(row=0, column=0, sticky="w")
        self.consultation_tree = ttk.Treeview(
            list_card,
            columns=("id", "title", "rows", "patterns"),
            show="headings",
            style="Data.Treeview",
        )
        for key, title, width in (
            ("id", "ID", 70),
            ("title", "Consulta / fecha", 320),
            ("rows", "Filas", 80),
            ("patterns", "Lote", 260),
        ):
            self.consultation_tree.heading(key, text=title)
            self.consultation_tree.column(key, width=width, anchor="center" if key in {"id", "rows"} else "w")
        consultation_tree_scroll = ttk.Scrollbar(list_card, orient=tk.VERTICAL, command=self.consultation_tree.yview)
        self.consultation_tree.configure(yscrollcommand=consultation_tree_scroll.set)
        self.consultation_tree.grid(row=1, column=0, sticky="nsew", pady=(12, 0))
        consultation_tree_scroll.grid(row=1, column=1, sticky="ns", pady=(12, 0))

        detail_card = ttk.Frame(body, style="Card.TFrame", padding=16)
        detail_card.grid(row=0, column=1, sticky="nsew")
        detail_card.columnconfigure(0, weight=1)
        detail_card.rowconfigure(1, weight=1)
        ttk.Label(detail_card, text="Detalle de la consulta", style="CardTitle.TLabel").grid(row=0, column=0, sticky="w")
        self.consultation_detail_text = self._create_readonly_text(detail_card)
        self.consultation_detail_text.grid(row=1, column=0, sticky="nsew", pady=(12, 0))

    def _build_pattern_tab(self) -> None:
        self.pattern_tab.columnconfigure(0, weight=3)
        self.pattern_tab.columnconfigure(1, weight=4)
        self.pattern_tab.rowconfigure(0, weight=1)

        pattern_list_card = ttk.Frame(self.pattern_tab, style="Card.TFrame", padding=16)
        pattern_list_card.grid(row=0, column=0, sticky="nsew", padx=(0, 8))
        pattern_list_card.columnconfigure(0, weight=1)
        pattern_list_card.rowconfigure(1, weight=1)
        ttk.Label(pattern_list_card, text="Mapa de patrones del proyecto", style="CardTitle.TLabel").grid(
            row=0, column=0, sticky="w"
        )
        self.pattern_tree = ttk.Treeview(
            pattern_list_card,
            columns=("pattern", "module"),
            show="headings",
            style="Data.Treeview",
        )
        self.pattern_tree.heading("pattern", text="Patron")
        self.pattern_tree.heading("module", text="Modulo")
        self.pattern_tree.column("pattern", width=150, anchor="center")
        self.pattern_tree.column("module", width=260, anchor="w")
        pattern_tree_scroll = ttk.Scrollbar(pattern_list_card, orient=tk.VERTICAL, command=self.pattern_tree.yview)
        self.pattern_tree.configure(yscrollcommand=pattern_tree_scroll.set)
        self.pattern_tree.grid(row=1, column=0, sticky="nsew", pady=(12, 0))
        pattern_tree_scroll.grid(row=1, column=1, sticky="ns", pady=(12, 0))

        detail_card = ttk.Frame(self.pattern_tab, style="Card.TFrame", padding=16)
        detail_card.grid(row=0, column=1, sticky="nsew")
        detail_card.columnconfigure(0, weight=1)
        detail_card.rowconfigure(1, weight=1)
        ttk.Label(detail_card, text="Explicacion aplicada", style="CardTitle.TLabel").grid(row=0, column=0, sticky="w")
        self.pattern_detail_text = self._create_readonly_text(detail_card)
        self.pattern_detail_text.grid(row=1, column=0, sticky="nsew", pady=(12, 0))

    def _bind_events(self) -> None:
        self.template_combo.bind("<<ComboboxSelected>>", lambda _event: self._load_selected_template())
        self.consultation_tree.bind("<<TreeviewSelect>>", lambda _event: self._render_selected_consultation())
        self.pattern_tree.bind("<<TreeviewSelect>>", lambda _event: self._render_selected_pattern())

    def _load_demo_defaults(self) -> None:
        self._apply_demo_credentials(self.service.demo_credentials()[0], auto_login=False)
        self._reset_order_form()
        self._reset_filters()
        self._sync_templates()
        self._render_patterns()

    def _apply_demo_credentials(self, data: dict[str, str], auto_login: bool = False) -> None:
        self.username_var.set(data["username"])
        self.password_var.set(data["password"])
        self.status_var.set(f"Credenciales demo cargadas para {data['role']}.")
        if auto_login:
            self._login()

    def _lock_workspace(self) -> None:
        for index in range(self.notebook.index("end")):
            self.notebook.tab(index, state="disabled")

    def _unlock_workspace(self) -> None:
        for index in range(self.notebook.index("end")):
            self.notebook.tab(index, state="normal")

    def _login(self) -> None:
        try:
            user = self.service.authenticate(self.username_var.get().strip(), self.password_var.get().strip())
        except Exception as exc:
            messagebox.showerror("MES Software", str(exc))
            self.status_var.set("No fue posible iniciar sesion.")
            return

        self.current_user = user
        self.session_var.set(
            f"Sesion activa | {user.username} | rol={user.role} | lineas={', '.join(user.allowed_lines) or 'ALL'}"
        )
        self.status_var.set("Sesion iniciada correctamente. El entorno software esta habilitado.")
        self.operator_id_var.set(self.operator_id_var.get().strip() or user.username.upper())
        self.operator_name_var.set(self.operator_name_var.get().strip() or user.username.replace("_", " ").title())
        if user.role != "admin":
            self.filter_requested_by_var.set(user.username)
        self._unlock_workspace()
        self._refresh_all_views()

    def _logout(self) -> None:
        self.current_user = None
        self.session_var.set("Sesion no iniciada")
        self.status_var.set("Sesion cerrada. Inicia sesion para continuar.")
        self.current_result = None
        self.current_consultations = []
        self._lock_workspace()
        self._clear_text(self.summary_text, "Inicia sesion para visualizar resultados.")
        self._clear_text(self.detail_text, "")
        self._clear_text(self.history_text, "")
        self._clear_text(self.security_text, "")
        self._clear_text(self.consultation_detail_text, "")
        self._set_tree_rows(self.recent_orders_tree, [])
        self._set_tree_rows(self.consultation_tree, [])
        self._clear_text(self.audit_text, "")

    def _refresh_all_views(self) -> None:
        if not self.current_user:
            return
        self._sync_templates()
        self._refresh_dashboard()
        self._load_persisted_consultation_history()
        self._render_patterns()

    def _refresh_dashboard(self) -> None:
        if not self.current_user:
            return
        snapshot = self.service.dashboard_snapshot(self.current_user)
        self.metric_vars["total_orders"].set(str(snapshot.get("total_orders", 0)))
        self.metric_vars["planned_units_total"].set(str(snapshot.get("planned_units_total", 0)))
        self.metric_vars["produced_units_total"].set(str(snapshot.get("produced_units_total", 0)))
        self.metric_vars["downtime_minutes_total"].set(str(snapshot.get("downtime_minutes_total", 0)))
        self.metric_vars["average_oee"].set(f"{snapshot.get('average_oee', 0.0)}%")
        self.metric_vars["total_consultation_batches"].set(str(snapshot.get("total_consultation_batches", 0)))

        rows = []
        for item in snapshot.get("recent_orders", []):
            shift_window = "-"
            if item.get("shift_start_at") and item.get("shift_end_at"):
                shift_window = (
                    f"{str(item.get('shift_start_at'))[11:16]} - "
                    f"{str(item.get('shift_end_at'))[11:16]}"
                )
            created_at = str(item.get("created_at", "-")).replace("T", " ")
            rows.append(
                (
                    item.get("order_id", ""),
                    item.get("production_line", ""),
                    item.get("shift", ""),
                    shift_window,
                    created_at,
                    item.get("state", ""),
                    f"{item.get('produced_units', 0)}/{item.get('planned_units', 0)}",
                )
            )
        self._set_tree_rows(self.recent_orders_tree, rows)

        audit_lines = []
        consultation_batches = snapshot.get("recent_consultation_batches", [])
        if consultation_batches:
            audit_lines.append("Historial reciente de consultas")
            audit_lines.append("")
            for batch in consultation_batches:
                audit_lines.append(
                    f"[{batch.get('created_at', '-')}] {batch.get('batch_code', '-')}\n"
                    f"user={batch.get('username', '-')}, queries={batch.get('total_queries', 0)}, "
                    f"rows={batch.get('total_rows', 0)}\n"
                    f"filters={batch.get('filters_json', '{}')}\n"
                )
            audit_lines.append("")
            audit_lines.append("Auditoria reciente")
            audit_lines.append("")
        for event in snapshot.get("recent_audits", []):
            audit_lines.append(
                f"[{event.get('created_at', '-')}] {event.get('event_type', '-')}\n"
                f"user={event.get('username', '-')}, order={event.get('order_id', '-')}, "
                f"line={event.get('production_line', '-')}\n"
                f"{event.get('description', '-')}\n"
            )
        self._clear_text(self.audit_text, "\n".join(audit_lines) if audit_lines else "Sin auditoria registrada.")

    def _sync_templates(self) -> None:
        template_keys = self.service.list_templates()
        self.template_combo.configure(values=template_keys)
        if not self.template_var.get() and template_keys:
            self.template_var.set(template_keys[0])
            self._load_selected_template()

    def _seed_demo_order(self) -> None:
        template_key = self.template_var.get().strip() or "cnc_standard"
        now_token = datetime.now().strftime("%Y%m%d%H%M%S")
        self.order_id_var.set(f"OP-{now_token}")
        self.units_var.set("320")
        if template_key.startswith("robot"):
            self.protocol_var.set("modbus")
            self.shift_var.set("NIGHT")
        else:
            self.protocol_var.set("opcua")
            self.shift_var.set("DAY")
        self.status_var.set("Se cargaron datos demo para una nueva ejecucion.")

    def _reset_order_form(self) -> None:
        self.order_form_scroll.reset_view()
        self.template_key_var.set("")
        self.template_name_var.set("")
        self.machine_type_var.set("cnc")
        self.line_var.set("")
        self.program_var.set("")
        self.product_code_var.set("")
        self.product_name_var.set("")
        self.line_mode_var.set("automatic")
        self.quality_checks_var.set("")
        self.order_id_var.set(f"OP-{datetime.now().strftime('%Y%m%d%H%M%S')}")
        self.units_var.set("300")
        self.shift_var.set("DAY")
        self.protocol_var.set("opcua")
        self.parameters_text.delete("1.0", tk.END)
        self.parameters_text.insert("1.0", '{\n  "speed_rpm": 1800\n}')
        if self.current_user:
            self.operator_id_var.set(self.current_user.username.upper())
            self.operator_name_var.set(self.current_user.username.replace("_", " ").title())
        else:
            self.operator_id_var.set("")
            self.operator_name_var.set("")
        self.current_result = None
        self._clear_text(self.summary_text, "Ejecuta una orden para ver su resumen operacional.")
        self._clear_text(self.detail_text, "")
        self._clear_text(self.history_text, "")
        self._clear_text(self.security_text, "")

    def _load_selected_template(self) -> None:
        template_key = self.template_var.get().strip()
        if not template_key:
            return
        try:
            data = self.service.get_template_definition(template_key)
        except Exception as exc:
            messagebox.showerror("MES Software", str(exc))
            return

        self.template_key_var.set(data["template_key"])
        self.template_name_var.set(data["template_name"])
        self.machine_type_var.set(data["machine_type"])
        self.line_var.set(data["line"])
        self.program_var.set(data["program"])
        self.product_code_var.set(data["product_code"])
        self.product_name_var.set(data["product_name"])
        self.line_mode_var.set(data["line_mode"])
        self.quality_checks_var.set(", ".join(data["quality_checks"]))
        self.parameters_text.delete("1.0", tk.END)
        self.parameters_text.insert("1.0", json.dumps(data["parameters"], ensure_ascii=False, indent=2))
        self.status_var.set(f"Plantilla {template_key} cargada correctamente.")

    def _save_template(self) -> None:
        try:
            self.service.register_template(
                template_key=self.template_key_var.get().strip(),
                template_name=self.template_name_var.get().strip(),
                machine_type=self.machine_type_var.get().strip(),
                line=self.line_var.get().strip(),
                program=self.program_var.get().strip(),
                product_code=self.product_code_var.get().strip(),
                product_name=self.product_name_var.get().strip(),
                quality_checks=self._parse_quality_checks(),
                line_mode=self.line_mode_var.get().strip(),
                parameters=self._parse_parameters(),
            )
        except Exception as exc:
            messagebox.showerror("MES Software", str(exc))
            self.status_var.set("No fue posible guardar la plantilla.")
            return

        self._sync_templates()
        self.template_var.set(self.template_key_var.get().strip())
        self.status_var.set(f"Plantilla {self.template_key_var.get().strip()} guardada correctamente.")

    def _parse_quality_checks(self) -> list[str]:
        return [item.strip() for item in self.quality_checks_var.get().split(",") if item.strip()]

    def _parse_parameters(self) -> dict:
        raw = self.parameters_text.get("1.0", tk.END).strip()
        if not raw:
            return {}
        parsed = json.loads(raw)
        if not isinstance(parsed, dict):
            raise ValueError("Los parametros deben venir en formato JSON objeto.")
        return parsed

    def _execute_order(self) -> None:
        if not self.current_user:
            messagebox.showwarning("MES Software", "Primero debes iniciar sesion.")
            return

        try:
            units = int(self.units_var.get().strip())
            request = OrderExecutionRequest(
                template_key=self.template_var.get().strip(),
                order_id=self.order_id_var.get().strip(),
                planned_units=units,
                shift=self.shift_var.get().strip(),
                protocol=self.protocol_var.get().strip(),
                operator_id=self.operator_id_var.get().strip(),
                operator_name=self.operator_name_var.get().strip(),
            )
            result = self.service.execute_order(self.current_user, request)
        except Exception as exc:
            messagebox.showerror("MES Software", str(exc))
            self.status_var.set("La orden no pudo ejecutarse.")
            return

        self.current_result = result
        self._render_result(result)
        self._refresh_dashboard()
        self.filter_order_id_var.set(request.order_id)
        self.filter_product_code_var.set(result.get("product", {}).get("product_code", ""))
        self.filter_line_var.set(result.get("startup", {}).get("production_line", ""))
        self.status_var.set(f"Orden {request.order_id} ejecutada y persistida correctamente.")
        self.notebook.select(self.order_tab)

    def _render_result(self, result: dict) -> None:
        if result.get("status") != "OK":
            self._clear_text(self.summary_text, json.dumps(result, ensure_ascii=False, indent=2))
            return

        operator = result.get("operator", {})
        product = result.get("product", {})
        final_report = result.get("final_report", {})
        secure_oee = result.get("secure_oee_report", {})

        summary_lines = [
            f"Status: {result.get('status')}",
            f"Order ID: {result.get('order_id')}",
            f"Operator: {operator.get('operator_id', '-')} / {operator.get('operator_name', '-')}",
            f"Product: {product.get('product_code', '-')} / {product.get('product_name', '-')}",
            f"Dispatch: {result.get('dispatch', '-')}",
            f"Planning timestamp: {result.get('state_history', [{}])[0].get('changed_at', '-') if result.get('state_history') else '-'}",
            f"Shift window: {result.get('persistence', {}).get('shift_start_at', '-')} -> {result.get('persistence', {}).get('shift_end_at', '-')}",
            f"Plant capacity units: {result.get('plant_capacity_units', '-')}",
            f"Stored DB: {result.get('persistence', {}).get('db_path', '-')}",
            f"Stored state: {result.get('persistence', {}).get('stored_state', '-')}",
            f"OEE: {secure_oee.get('oee', '-')}% ({secure_oee.get('source', '-')})",
        ]
        self._clear_text(self.summary_text, "\n".join(summary_lines))

        operation_lines = [
            "Final report",
            json.dumps(final_report, ensure_ascii=False, indent=2),
            "",
            "Execution log",
        ]
        for idx, item in enumerate(result.get("execution_log", []), start=1):
            operation_lines.append(
                f"{idx:02d}. {item.get('step', '')} | {item.get('pattern', '')} | {item.get('evidence', '')}"
            )
        self._clear_text(self.detail_text, "\n".join(operation_lines))

        history_lines = ["State history"]
        for item in result.get("state_history", []):
            history_lines.append(
                f"{item.get('changed_at')} | {item.get('from_state')} -> {item.get('to_state')} | {item.get('note')}"
            )
        history_lines.append("")
        history_lines.append("Observer notifications")
        for item in result.get("observer_notifications", []):
            history_lines.append(f"- {item}")
        self._clear_text(self.history_text, "\n".join(history_lines))

        security_lines = [
            "Secure OEE report",
            json.dumps(result.get("secure_oee_report", {}), ensure_ascii=False, indent=2),
            "",
            "Proxy audit log",
            json.dumps(result.get("proxy_audit_log", []), ensure_ascii=False, indent=2),
        ]
        self._clear_text(self.security_text, "\n".join(security_lines))

    def _build_filters(self) -> QueryFilters:
        limit = int(self.filter_limit_var.get().strip() or "25")
        requested_by = self.filter_requested_by_var.get().strip()
        if self.current_user and self.current_user.role != "admin":
            requested_by = self.current_user.username
            self.filter_requested_by_var.set(requested_by)

        return QueryFilters(
            order_id=self.filter_order_id_var.get().strip() or None,
            product_code=self.filter_product_code_var.get().strip() or None,
            production_line=self.filter_line_var.get().strip() or None,
            shift=self.filter_shift_var.get().strip() or None,
            protocol=self.filter_protocol_var.get().strip() or None,
            operator_id=self.filter_operator_id_var.get().strip() or None,
            state=self.filter_state_var.get().strip() or None,
            requested_by=requested_by or None,
            date_from=self.filter_date_from_var.get().strip() or None,
            date_to=self.filter_date_to_var.get().strip() or None,
            hour_from=self.filter_hour_from_var.get().strip() or None,
            hour_to=self.filter_hour_to_var.get().strip() or None,
            limit=limit,
        )

    def _generate_consultations(self) -> None:
        if not self.current_user:
            messagebox.showwarning("MES Software", "Primero debes iniciar sesion.")
            return

        try:
            filters = self._build_filters()
            self.current_consultations = self.service.generate_filtered_consultations(self.current_user, filters)
        except Exception as exc:
            messagebox.showerror("MES Software", str(exc))
            self.status_var.set("No fue posible generar las consultas.")
            return

        self._populate_consultation_tree()
        self.status_var.set(
            f"Consultas filtradas generadas correctamente: {len(self.current_consultations)} resultados SQ registrados."
        )
        self.notebook.select(self.consultation_tab)

    def _load_persisted_consultation_history(self) -> None:
        if not self.current_user:
            return

        try:
            limit = int(self.filter_limit_var.get().strip() or "25")
        except ValueError:
            limit = 25

        try:
            self.current_consultations = self.service.load_persisted_consultations(self.current_user, limit=max(limit, 5))
        except Exception as exc:
            messagebox.showerror("MES Software", str(exc))
            self.status_var.set("No fue posible recargar el historial persistido.")
            return

        self._populate_consultation_tree()

    def _populate_consultation_tree(self) -> None:
        rows = []
        for item in self.current_consultations:
            rows.append(
                (
                    item["identifier"],
                    f"{item['title']} | {item.get('generated_at', '-')}",
                    item["row_count"],
                    f"{item.get('batch_code', '-')}",
                )
            )
        self._set_tree_rows(self.consultation_tree, rows)
        if rows:
            first = self.consultation_tree.get_children()[0]
            self.consultation_tree.selection_set(first)
            self._render_selected_consultation()
        else:
            self._clear_text(
                self.consultation_detail_text,
                "No hay consultas persistidas para el usuario actual en la base de datos activa.",
            )

    def _render_selected_consultation(self) -> None:
        selection = self.consultation_tree.selection()
        if not selection:
            return
        item = self.consultation_tree.item(selection[0], "values")
        consultation_id = item[0]
        batch_code = item[3] if len(item) > 3 else ""
        consultation = next(
            (
                row for row in self.current_consultations
                if row["identifier"] == consultation_id and row.get("batch_code", "") == batch_code
            ),
            None,
        )
        if not consultation:
            return

        lines = [
            f"{consultation['identifier']} | {consultation['title']}",
            "",
            f"Lote: {consultation.get('batch_code', '-')}",
            f"Generada: {consultation.get('generated_at', '-')}",
            f"Usuario: {consultation.get('generated_by', self.current_user.username if self.current_user else '-')}",
            f"Objetivo: {consultation['objective']}",
            f"Pregunta: {consultation['business_question']}",
            f"Filas encontradas: {consultation['row_count']}",
            "",
            "Filtros aplicados:",
        ]
        for row in consultation["filters"]:
            lines.append(f"- {row}")
        lines.extend(
            [
                "",
                "SQL generado:",
                consultation["suggested_query"],
                "",
                f"Parametros: {consultation['query_params']}",
                "",
                "Patrones aplicados:",
            ]
        )
        for pattern in consultation["pattern_details"]:
            lines.append(f"- {pattern['pattern']}: {pattern['objective']}")
        lines.extend(["", "Muestra de resultados:", json.dumps(consultation["sample_result"], ensure_ascii=False, indent=2)])
        self._clear_text(self.consultation_detail_text, "\n".join(lines))

    def _export_consultations_doc(self) -> None:
        if not self.current_consultations:
            messagebox.showwarning("MES Software", "Primero genera las consultas filtradas.")
            return

        initial_dir = Path(__file__).resolve().parent
        target_path = filedialog.asksaveasfilename(
            title="Guardar documento de consultas MES",
            defaultextension=".docx",
            initialdir=str(initial_dir),
            initialfile=f"consultas_mes_software_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
            filetypes=[("Word Document", "*.docx")],
        )
        if not target_path:
            return

        document = Document()
        document.add_heading("MES Software Control Center - Consultas filtradas", level=0)
        document.add_paragraph(
            "Documento exportado desde la interfaz profesional de la capa software del proyecto MES."
        )

        for consultation in self.current_consultations:
            document.add_heading(f"{consultation['identifier']} - {consultation['title']}", level=1)
            document.add_paragraph(f"Objetivo: {consultation['objective']}")
            document.add_paragraph(f"Pregunta: {consultation['business_question']}")
            document.add_paragraph("Filtros aplicados:")
            for row in consultation["filters"]:
                document.add_paragraph(row, style="List Bullet")
            document.add_paragraph("SQL generado:")
            document.add_paragraph(consultation["suggested_query"])
            document.add_paragraph(f"Parametros: {consultation['query_params']}")
            document.add_paragraph("Patrones aplicados:")
            for pattern in consultation["pattern_details"]:
                document.add_paragraph(
                    f"{pattern['pattern']} - {pattern['objective']} ({pattern['usage_point']})",
                    style="List Bullet",
                )
            document.add_paragraph("Muestra de resultados:")
            document.add_paragraph(json.dumps(consultation["sample_result"], ensure_ascii=False, indent=2))

        document.save(target_path)
        self.status_var.set(f"Documento exportado: {Path(target_path).name}")
        messagebox.showinfo("MES Software", f"Documento generado correctamente:\n{target_path}")

    def _reset_filters(self) -> None:
        today = datetime.now().date().isoformat()
        self.filter_order_id_var.set("")
        self.filter_product_code_var.set("")
        self.filter_line_var.set("")
        self.filter_shift_var.set("")
        self.filter_protocol_var.set("")
        self.filter_operator_id_var.set("")
        self.filter_state_var.set("")
        self.filter_requested_by_var.set(self.current_user.username if self.current_user and self.current_user.role != "admin" else "")
        self.filter_date_from_var.set(today)
        self.filter_date_to_var.set(today)
        self.filter_hour_from_var.set("00:00:00")
        self.filter_hour_to_var.set("23:59:59")
        self.filter_limit_var.set("25")
        self.current_consultations = []
        self._set_tree_rows(self.consultation_tree, [])
        self._clear_text(self.consultation_detail_text, "Configura filtros y genera las consultas persistentes.")

    def _render_patterns(self) -> None:
        rows = [(item["pattern"], item["module"]) for item in self.pattern_map]
        self._set_tree_rows(self.pattern_tree, rows)
        if rows:
            first = self.pattern_tree.get_children()[0]
            self.pattern_tree.selection_set(first)
            self._render_selected_pattern()

    def _render_selected_pattern(self) -> None:
        selection = self.pattern_tree.selection()
        if not selection:
            return
        values = self.pattern_tree.item(selection[0], "values")
        name = values[0]
        item = next((row for row in self.pattern_map if row["pattern"] == name), None)
        if not item:
            return

        lines = [
            f"Patron: {item['pattern']}",
            "",
            f"Objetivo: {item['objective']}",
            f"Modulo: {item['module']}",
            f"Clases clave: {item['key_classes']}",
            f"Punto de uso: {item['usage_point']}",
        ]
        self._clear_text(self.pattern_detail_text, "\n".join(lines))

    def _labeled_entry(
        self,
        parent: ttk.Widget,
        label: str,
        variable: tk.StringVar,
        row: int,
        column: int,
        colspan: int = 1,
    ) -> None:
        ttk.Label(parent, text=label).grid(row=row, column=column, sticky="w", padx=(0, 8), pady=(6, 4))
        ttk.Entry(parent, textvariable=variable).grid(
            row=row + 1, column=column, columnspan=colspan, sticky="ew", padx=(0, 8), pady=(0, 8)
        )

    def _labeled_combo(
        self,
        parent: ttk.Widget,
        label: str,
        variable: tk.StringVar,
        values: tuple[str, ...],
        row: int,
        column: int,
    ) -> None:
        ttk.Label(parent, text=label).grid(row=row, column=column, sticky="w", padx=(0, 8), pady=(6, 4))
        ttk.Combobox(parent, textvariable=variable, values=values, state="readonly").grid(
            row=row + 1, column=column, sticky="ew", padx=(0, 8), pady=(0, 8)
        )

    def _create_readonly_text(self, parent: ttk.Widget) -> ScrolledText:
        widget = ScrolledText(parent, wrap=tk.WORD, font=("Consolas", 10), relief=tk.FLAT)
        widget.configure(state=tk.DISABLED, background="#ffffff", foreground="#0f172a")
        return widget

    def _clear_text(self, widget: ScrolledText, content: str) -> None:
        widget.configure(state=tk.NORMAL)
        widget.delete("1.0", tk.END)
        widget.insert(tk.END, content)
        widget.configure(state=tk.DISABLED)

    def _set_tree_rows(self, tree: ttk.Treeview, rows: list[tuple]) -> None:
        for item in tree.get_children():
            tree.delete(item)
        for row in rows:
            tree.insert("", tk.END, values=row)


def main() -> None:
    app = MESSoftwareControlCenter()
    app.mainloop()


if __name__ == "__main__":
    main()
