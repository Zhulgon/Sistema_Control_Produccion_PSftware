from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt


OUTPUT_PATH = (
    Path(__file__).resolve().parents[1]
    / "Documentos 3 corte"
    / "DOCUMENTOS"
    / "Informe_Tercer_Corte_MES_Software.docx"
)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)

    for style_name, size in (("Heading 1", 14), ("Heading 2", 12)):
        style = document.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True


def add_paragraph(document: Document, text: str, *, bold: bool = False, center: bool = False) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def add_bullet(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.line_spacing = 1.5
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def add_heading(document: Document, text: str, level: int = 1) -> None:
    heading = document.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)

    for row_values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_values):
            cells[idx].text = value
            for run in cells[idx].paragraphs[0].runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(10)


def build_document() -> None:
    document = Document()
    configure_document(document)

    for _ in range(4):
        add_paragraph(document, "", center=True)
    add_paragraph(document, "DOCUMENTO TÉCNICO DEL TERCER CORTE", bold=True, center=True)
    add_paragraph(
        document,
        "Escalamiento del Sistema MES de Control de Producción hacia una solución de software",
        bold=True,
        center=True,
    )
    for _ in range(3):
        add_paragraph(document, "", center=True)
    add_paragraph(document, "Asignatura: Patrones de Software", center=True)
    add_paragraph(document, "Programa: Ingeniería de Sistemas", center=True)
    add_paragraph(
        document,
        "Integrantes: Juan Sebastian Rubiano Romero - Angie Lucero Vargas Amado",
        center=True,
    )
    add_paragraph(document, "Proyecto: Sistema de Control de Producción (MES)", center=True)
    add_paragraph(document, f"Fecha de elaboración: {datetime.now().strftime('%Y-%m-%d')}", center=True)
    document.add_page_break()
    add_heading(document, "1. Introduccion")
    add_paragraph(
        document,
        "El presente documento describe la evolucion del Sistema de Control de Produccion (MES) hacia una propuesta mas cercana a un producto de software utilizable, defendible y con valor para un posible comprador. El proyecto deja de centrarse solo en la demostracion academica de patrones y pasa a mostrar como una solucion MES puede registrar operaciones, proteger acceso, conservar evidencia y apoyar decisiones reales de planta.",
    )
    add_paragraph(
        document,
        "La propuesta integra persistencia, seguridad minima, trazabilidad operativa y un paquete de 10 consultas de negocio orientadas a capacidad, cumplimiento, integracion, desempeno, portafolio de productos y control operativo. Esto permite sustentar no solo como esta construido el sistema, sino por que seria util adoptarlo dentro de un contexto productivo.",
    )
    add_heading(document, "2. Justificacion del escalamiento")
    add_paragraph(
        document,
        "Para que el proyecto tenga valor ante una persona interesada en adquirir o adoptar la solucion, no basta con ejecutar una orden de produccion de principio a fin. El sistema debe responder preguntas concretas sobre demanda, carga, capacidad, eficiencia, productos, turnos, operadores, trazabilidad y seguridad. Por ello se fortalecio la capa de persistencia, se controlaron accesos y se redefinio el bloque analitico para que las consultas entreguen informacion util para tomar decisiones y no solo evidencia academica.",
    )
    add_heading(document, "3. Alcance del proyecto")
    add_paragraph(
        document,
        "El alcance de esta fase comprende la transformacion del proyecto en una base de software funcional orientada a la gestion de ordenes MES, manteniendo la coherencia con los patrones de software trabajados en la asignatura y fortaleciendo el valor practico de la solucion.",
    )
    for item in [
        "Registrar, cargar y reutilizar plantillas de produccion para distintas lineas y tipos de maquina.",
        "Ejecutar ordenes de produccion con trazabilidad desde la planificacion hasta el cierre.",
        "Persistir usuarios, ordenes, historial de estados, telemetria, auditoria y lotes de consultas en SQLite.",
        "Aplicar autenticacion minima con proteccion de contrasenas mediante hash y salt.",
        "Restringir acceso a reportes y consultas segun rol y linea autorizada.",
        "Generar y almacenar 10 consultas de negocio filtrables por fecha, hora, turno, linea, producto, protocolo, usuario y estado.",
        "Presentar evidencia util para evaluar capacidad, cumplimiento, portafolio, seguridad, trazabilidad e integracion operativa.",
        "Integrar patrones creacionales, estructurales y de comportamiento en un mismo flujo de negocio.",
    ]:
        add_bullet(document, item)
    add_heading(document, "4. Fuera de alcance")
    for item in [
        "Despliegue del sistema en ambiente productivo web o móvil.",
        "Integración con PLC físicos, sensores reales o infraestructura industrial de planta.",
        "Implementación de mecanismos de seguridad avanzados como OAuth, JWT, MFA o cifrado empresarial.",
        "Arquitectura distribuida basada en microservicios o balanceo de carga real.",
        "Analítica predictiva, reportería BI o dashboards en tiempo real con herramientas externas.",
    ]:
        add_bullet(document, item)
    add_heading(document, "5. Objetivos")
    add_heading(document, "5.1 Objetivo general", level=2)
    add_paragraph(
        document,
        "Desarrollar la evolucion del Sistema de Control de Produccion MES hacia una solucion de software escalable, segura, trazable y orientada a la toma de decisiones, integrando persistencia, usuarios, consultas de negocio y patrones de software alineados con el flujo operativo del proceso productivo.",
    )
    add_heading(document, "5.2 Objetivos especificos", level=2)
    for item in [
        "Implementar una base de datos ligera para persistir la informacion operativa y analitica del sistema.",
        "Gestionar usuarios y permisos minimos para controlar el acceso a ordenes, lineas y reportes.",
        "Aplicar el patron State para modelar el ciclo de vida de la orden de produccion.",
        "Aplicar el patron Observer para notificar y auditar cambios de estado sin acoplar la logica principal.",
        "Integrar una interfaz grafica profesional que unifique autenticacion, operacion, consultas y evidencias.",
        "Consolidar un portafolio de 10 consultas de negocio que permita demostrar valor operativo y comercial del sistema MES.",
        "Persistir lotes de consultas para fortalecer la trazabilidad analitica y la defensa del sistema como producto de software.",
        "Documentar el alcance, los resultados y el papel de cada patron dentro del proyecto.",
    ]:
        add_bullet(document, item)
    add_heading(document, "6. Descripción general de la solución")
    add_paragraph(
        document,
        "La solución construida se organiza en una capa de interfaz, una capa de aplicación, una capa de "
        "persistencia y un conjunto de módulos de patrones funcionales. La interfaz principal permite iniciar "
        "sesión, ejecutar órdenes, consultar resultados, visualizar patrones y exportar evidencia. La capa de "
        "servicio orquesta la ejecución del flujo MES, gestiona el acceso del usuario autenticado y coordina la "
        "persistencia. La base de datos SQLite almacena la información estructural y operativa del sistema. "
        "Finalmente, los módulos de patrones conservan la demostración arquitectónica del proyecto y la conectan "
        "con los requerimientos reales de la aplicación.",
    )
    add_heading(document, "7. Componentes principales del software")
    add_table(
        document,
        ["Componente", "Responsabilidad"],
        [
            ["UI_MES_Software_Pro.py", "Interfaz principal del sistema. Integra login, dashboard, operación, consultas y evidencias."],
            ["software_mes_service.py", "Orquestación de la capa software y coordinación del flujo MES con seguridad y persistencia."],
            ["software_mes_persistence.py", "Gestión de SQLite, creación de esquema y almacenamiento de entidades clave."],
            ["software_mes_security.py", "Autenticación, hash de contraseñas y control de acceso por rol y línea."],
            ["software_mes_queries.py", "Construcción de consultas filtrables por fecha, hora y contexto operativo."],
            ["software_mes_behavioral.py", "Implementación de patrones State y Observer para el tercer corte."],
            ["mes_functional_app.py", "Base funcional previa sobre la cual se mantiene el flujo de patrones del proyecto."],
        ],
    )
    add_heading(document, "8. Persistencia, base de datos y seguridad mínima")
    add_paragraph(
        document,
        "Con el fin de escalar el proyecto hacia una solución más defendible como software, se implementó una "
        "base de datos SQLite que centraliza la información del sistema. La persistencia ya no se limita a la "
        "ejecución en memoria, sino que registra órdenes, historial de estados, auditoría, telemetría y lotes "
        "de consultas.",
    )
    add_table(
        document,
        ["Tabla", "Propósito"],
        [
            ["users", "Almacenar credenciales, roles, líneas autorizadas y estado del usuario."],
            ["orders", "Persistir órdenes ejecutadas, producto, turno, protocolo, ventana horaria y estado final."],
            ["order_history", "Registrar las transiciones de estado de cada orden para mantener trazabilidad."],
            ["telemetry_events", "Guardar eventos asociados a telemetría y producción por orden."],
            ["audit_events", "Conservar evidencia de login, proxy, seguridad y cambios relevantes."],
            ["consultation_batches", "Registrar cada lote de consultas generado desde la interfaz."],
            ["consultation_results", "Persistir el detalle de cada consulta, su SQL, filtros y resultados."],
        ],
    )
    add_paragraph(
        document,
        "En materia de seguridad mínima, el sistema incorpora autenticación basada en usuario y contraseña, "
        "protección de contraseñas con hash y salt, roles diferenciados y restricción de acceso por línea de "
        "producción. Este nivel de seguridad no pretende cubrir un entorno empresarial completo, pero sí "
        "constituye un avance coherente respecto a una práctica puramente teórica.",
    )
    add_heading(document, "9. Consultas de negocio y mejora analitica")
    add_paragraph(
        document,
        "El bloque analitico del sistema fue reorganizado para responder preguntas que realmente ayuden a justificar la adopcion del MES. Las consultas ya no se limitan a mostrar datos operativos sueltos: ahora permiten explicar demanda, integracion tecnica, presion de capacidad, cumplimiento, trazabilidad, productividad, balance de lineas, seguridad, comportamiento del portafolio y diferencias entre turnos. Adicionalmente, cada lote de consultas queda persistido con filtros, parametros y muestra de resultados, lo que mejora la trazabilidad de la capa analitica.",
    )
    add_table(
        document,
        ["ID", "Consulta", "Utilidad"],
        [
            ["SQ1", "Demanda confirmada por orden y contexto", "Identificar que ordenes, productos y lineas concentraron la demanda registrada."],
            ["SQ2", "Integracion operativa por protocolo", "Evidenciar soporte multi-protocolo y facilidad de integracion con entornos industriales."],
            ["SQ3", "Capacidad disponible vs carga comprometida", "Mostrar si la demanda presiona la capacidad instalada y donde conviene balancear recursos."],
            ["SQ4", "Cumplimiento de produccion y OEE", "Medir cumplimiento del plan, eficiencia real y tiempos de parada por orden."],
            ["SQ5", "Trazabilidad y evidencia de cumplimiento", "Demostrar seguimiento end-to-end y evidencia util para calidad, auditoria y confianza operativa."],
            ["SQ6", "Productividad por operador", "Comparar desempeno humano y detectar oportunidades de mejora operacional."],
            ["SQ7", "Carga y utilizacion por linea", "Balancear lineas, priorizar recursos y detectar concentraciones de trabajo."],
            ["SQ8", "Seguridad operativa y auditoria", "Probar control de accesos, rastreabilidad de acciones y gobierno operativo."],
            ["SQ9", "Desempeno por producto", "Priorizar productos por volumen, cumplimiento y eficiencia para soportar decisiones de portafolio."],
            ["SQ10", "Comparativo de desempeno por turno", "Detectar diferencias entre DAY y NIGHT para intervenir la operacion con mejor criterio."],
        ],
    )
    add_heading(document, "10. Patrones aplicados dentro del proyecto")
    add_table(
        document,
        ["Patrón", "Categoría", "Aplicación en el proyecto", "Aporte principal"],
        [
            ["Singleton", "Creacional", "Centralizar el estado global del MES y la instancia de base de datos.", "Evitar duplicidad de control y mantener consistencia."],
            ["Factory Method", "Creacional", "Crear máquinas concretas según el tipo de orden.", "Reducir acoplamiento entre cliente y clases concretas."],
            ["Abstract Factory", "Creacional", "Construir familias compatibles de línea, máquina y asistente.", "Garantizar coherencia entre componentes relacionados."],
            ["Builder", "Creacional", "Construir reportes y estructuras de consulta paso a paso.", "Mejorar claridad en objetos complejos."],
            ["Prototype", "Creacional", "Clonar plantillas de producción reutilizables.", "Acelerar el registro y disminuir configuración repetitiva."],
            ["Adapter", "Estructural", "Integrar la capa MES con un gateway legacy.", "Permitir compatibilidad con sistemas heredados."],
            ["Bridge", "Estructural", "Separar el tipo de orden del protocolo industrial.", "Permitir variar protocolos y órdenes sin romper el cliente."],
            ["Composite", "Estructural", "Modelar planta, líneas y estaciones como jerarquía.", "Calcular capacidad total de forma uniforme."],
            ["Decorator", "Estructural", "Enriquecer reportes con calidad, trazabilidad y OEE.", "Agregar responsabilidades sin modificar el objeto base."],
            ["Facade", "Estructural", "Simplificar el punto de entrada del flujo operativo.", "Disminuir complejidad para el cliente de la aplicación."],
            ["Flyweight", "Estructural", "Compartir perfiles de máquina en telemetría repetitiva.", "Ahorrar memoria y mejorar escalabilidad en eventos."],
            ["Proxy", "Estructural", "Proteger reportes sensibles con autorización, caché y auditoría.", "Controlar acceso y reforzar la seguridad mínima."],
            ["State", "Comportamiento", "Modelar las transiciones CREATED, PLANNED, DISPATCHED, RUNNING y COMPLETED.", "Evitar condicionales frágiles y mejorar mantenibilidad."],
            ["Observer", "Comportamiento", "Notificar y auditar cambios de estado de la orden.", "Desacoplar reacciones del flujo principal."],
        ],
    )
    add_heading(document, "11. Resultados obtenidos")
    for item in [
        "Se consolido una interfaz profesional para autenticacion, dashboard, operacion y consultas del sistema MES.",
        "Se implemento una base SQLite con persistencia de usuarios, ordenes, historial, telemetria, auditoria y lotes de consultas.",
        "Se incorporo seguridad minima basada en hash, salt, roles y lineas autorizadas.",
        "Se reorganizo el bloque analitico en 10 consultas de negocio con mayor utilidad para operacion, seguimiento y sustentacion comercial.",
        "Se almacenaron los lotes de consultas para fortalecer la trazabilidad analitica del sistema.",
        "Se integraron State y Observer como patrones de comportamiento alineados con las orientaciones del tercer corte.",
        "Se mantuvo la demostracion de los patrones previos sin perder la funcionalidad ya construida en cortes anteriores.",
        "Se verifico el flujo principal mediante pruebas end-to-end del escalamiento del sistema.",
    ]:
        add_bullet(document, item)
    add_heading(document, "12. Conclusiones")
    add_paragraph(
        document,
        "La evolucion del proyecto demuestra que los patrones de software aportan mas valor cuando se conectan con preguntas reales del negocio. En esta version, el MES no solo registra ordenes: tambien explica demanda, capacidad, cumplimiento, seguridad, trazabilidad y comportamiento del portafolio por medio de consultas defendibles desde una perspectiva operativa.",
    )
    add_paragraph(
        document,
        "Desde la mirada de una persona que evaluara adquirir la solucion, el mayor avance consiste en que el sistema ya puede mostrar utilidad concreta: ayuda a entender donde esta la carga, como rinden las lineas y los turnos, que productos concentran el esfuerzo y que evidencia existe para auditar lo ocurrido. Eso hace que la propuesta sea mas argumentable, mas cercana a un caso real y mas facil de presentar como software con valor." ,
    )
    add_heading(document, "13. Recomendaciones y trabajo futuro")
    for item in [
        "Extender la gestión de turnos para permitir configuración paramétrica de horarios desde la interfaz.",
        "Agregar reportes históricos más avanzados y comparativos entre líneas, productos y períodos.",
        "Integrar exportaciones adicionales a PDF o reportes tabulares automatizados.",
        "Explorar autenticación reforzada y bitácoras más detalladas si el proyecto evoluciona a un escenario institucional.",
        "Conectar la persistencia actual con servicios externos o APIs industriales en una fase posterior.",
    ]:
        add_bullet(document, item)
    add_heading(document, "14. Cierre")
    add_paragraph(
        document,
        "En conclusión, el Sistema MES desarrollado en la asignatura demuestra una transición clara desde la "
        "comprensión conceptual de los patrones hasta su aplicación coordinada en una solución de software. El "
        "tercer corte consolida ese avance al incorporar persistencia, seguridad, consultas filtradas y patrones "
        "de comportamiento, logrando un proyecto más robusto, más argumentable y más cercano a un caso real de "
        "ingeniería de software.",
    )

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(OUTPUT_PATH))
    print(f"Documento generado en: {OUTPUT_PATH}")


if __name__ == "__main__":
    build_document()
