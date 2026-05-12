"""
Mini interfaz de escritorio para el proyecto MES funcional.

Ejecucion:
python UI_MES_Proyecto.py
"""

from __future__ import annotations

import json
import tkinter as tk
from datetime import datetime
from pathlib import Path
from tkinter import filedialog, messagebox, ttk
from tkinter.scrolledtext import ScrolledText

from docx import Document

from mes_functional_app import OrderExecutionRequest, build_default_mes_project


class MESMiniUI(tk.Tk):
    def __init__(self) -> None:
        super().__init__()
        self.title("Sistema MES - Sustentacion del Proyecto")
        self.geometry("1280x820")
        self.minsize(1180, 720)

        self.app = build_default_mes_project()
        self.last_request: OrderExecutionRequest | None = None
        self.last_result: dict | None = None
        self.saved_consultations: list[dict] = []
        self.consultation_batch_label = ""
        self._placeholder_vars: dict[str, str] = {}
        self._placeholder_entries: list[tuple[ttk.Entry, tk.StringVar]] = []
        self._text_placeholder = '{\n  "speed_rpm": 1800\n}'
        self._text_placeholder_active = False
        self._build_theme()
        self._build_layout()
        self._sync_template_options()
        self._reset_form()

    def _build_theme(self) -> None:
        self.palette = {
            "bg": "#e9eef5",
            "surface": "#f8fafc",
            "surface_alt": "#ffffff",
            "border": "#cbd5e1",
            "border_strong": "#94a3b8",
            "text": "#0f172a",
            "muted": "#475569",
            "hint": "#64748b",
            "primary": "#0f4c81",
            "primary_hover": "#0b3a63",
            "primary_active": "#082a47",
            "accent": "#0f766e",
            "accent_hover": "#115e59",
            "tab_idle": "#dbe4ef",
            "tab_active": "#ffffff",
            "input_bg": "#ffffff",
            "input_fg": "#0f172a",
            "placeholder": "#94a3b8",
            "text_panel": "#f7fafc",
        }

        self.configure(background=self.palette["bg"])
        style = ttk.Style(self)
        style.theme_use("clam")

        style.configure(
            ".",
            background=self.palette["bg"],
            foreground=self.palette["text"],
            font=("Segoe UI", 10),
        )
        style.configure("App.TFrame", background=self.palette["bg"])
        style.configure("Surface.TFrame", background=self.palette["surface"])
        style.configure(
            "Card.TFrame",
            background=self.palette["surface"],
            borderwidth=1,
            relief="solid",
        )
        style.configure(
            "Header.TLabel",
            font=("Segoe UI Semibold", 20),
            foreground=self.palette["text"],
            background=self.palette["bg"],
        )
        style.configure(
            "Subheader.TLabel",
            font=("Segoe UI", 10),
            foreground=self.palette["muted"],
            background=self.palette["bg"],
        )
        style.configure(
            "Hint.TLabel",
            font=("Segoe UI", 10),
            foreground=self.palette["hint"],
            background=self.palette["surface"],
        )
        style.configure(
            "Status.TLabel",
            font=("Segoe UI Semibold", 10),
            foreground=self.palette["primary"],
            background=self.palette["surface"],
        )
        style.configure(
            "Field.TLabel",
            font=("Segoe UI Semibold", 9),
            foreground=self.palette["muted"],
            background=self.palette["surface"],
        )
        style.configure(
            "Card.TLabelframe",
            background=self.palette["surface"],
            bordercolor=self.palette["border"],
            borderwidth=1,
            relief="solid",
            padding=14,
        )
        style.configure(
            "Card.TLabelframe.Label",
            background=self.palette["surface"],
            foreground=self.palette["text"],
            font=("Segoe UI Semibold", 11),
        )
        style.configure(
            "TButton",
            background=self.palette["surface_alt"],
            foreground=self.palette["text"],
            bordercolor=self.palette["border"],
            borderwidth=1,
            relief="flat",
            focusthickness=0,
            focuscolor=self.palette["surface_alt"],
            padding=(14, 9),
            font=("Segoe UI Semibold", 9),
        )
        style.map(
            "TButton",
            background=[
                ("active", self.palette["surface_alt"]),
                ("pressed", "#eef2f7"),
            ],
            bordercolor=[
                ("active", self.palette["border_strong"]),
                ("pressed", self.palette["border_strong"]),
            ],
        )
        style.configure(
            "Primary.TButton",
            background=self.palette["primary"],
            foreground="#ffffff",
            bordercolor=self.palette["primary"],
            padding=(16, 10),
            font=("Segoe UI Semibold", 10),
        )
        style.map(
            "Primary.TButton",
            background=[
                ("active", self.palette["primary_hover"]),
                ("pressed", self.palette["primary_active"]),
            ],
            bordercolor=[
                ("active", self.palette["primary_hover"]),
                ("pressed", self.palette["primary_active"]),
            ],
            foreground=[("disabled", "#d8e0eb")],
        )
        style.configure(
            "Accent.TButton",
            background=self.palette["accent"],
            foreground="#ffffff",
            bordercolor=self.palette["accent"],
        )
        style.map(
            "Accent.TButton",
            background=[
                ("active", self.palette["accent_hover"]),
                ("pressed", self.palette["accent_hover"]),
            ],
            bordercolor=[
                ("active", self.palette["accent_hover"]),
                ("pressed", self.palette["accent_hover"]),
            ],
        )
        style.configure(
            "TEntry",
            fieldbackground=self.palette["input_bg"],
            foreground=self.palette["input_fg"],
            bordercolor=self.palette["border"],
            lightcolor=self.palette["border"],
            darkcolor=self.palette["border"],
            insertcolor=self.palette["input_fg"],
            padding=(10, 8),
        )
        style.map(
            "TEntry",
            bordercolor=[("focus", self.palette["primary"])],
            lightcolor=[("focus", self.palette["primary"])],
            darkcolor=[("focus", self.palette["primary"])],
        )
        style.configure(
            "TCombobox",
            fieldbackground=self.palette["input_bg"],
            foreground=self.palette["input_fg"],
            bordercolor=self.palette["border"],
            lightcolor=self.palette["border"],
            darkcolor=self.palette["border"],
            arrowcolor=self.palette["primary"],
            padding=(8, 8),
        )
        style.map(
            "TCombobox",
            bordercolor=[("focus", self.palette["primary"])],
            lightcolor=[("focus", self.palette["primary"])],
            darkcolor=[("focus", self.palette["primary"])],
        )
        style.configure(
            "TSpinbox",
            fieldbackground=self.palette["input_bg"],
            foreground=self.palette["input_fg"],
            bordercolor=self.palette["border"],
            lightcolor=self.palette["border"],
            darkcolor=self.palette["border"],
            arrowcolor=self.palette["primary"],
            padding=(8, 8),
        )
        style.configure(
            "TNotebook",
            background=self.palette["bg"],
            borderwidth=0,
            tabmargins=(0, 0, 0, 0),
        )
        style.configure(
            "TNotebook.Tab",
            background=self.palette["tab_idle"],
            foreground=self.palette["muted"],
            padding=(18, 10),
            font=("Segoe UI Semibold", 10),
            borderwidth=0,
        )
        style.map(
            "TNotebook.Tab",
            background=[("selected", self.palette["tab_active"])],
            foreground=[("selected", self.palette["primary"])],
        )

    def _build_layout(self) -> None:
        root = ttk.Frame(self, style="App.TFrame", padding=18)
        root.pack(fill=tk.BOTH, expand=True)

        title = ttk.Label(root, text="Sistema MES - Evidencias del Proyecto", style="Header.TLabel")
        title.pack(anchor="w")
        subtitle = ttk.Label(
            root,
            text="Registra la orden, ejecuta el flujo de produccion y presenta los resultados del proyecto.",
            style="Subheader.TLabel",
        )
        subtitle.pack(anchor="w", pady=(4, 16))

        top = ttk.LabelFrame(root, text="Registro de la orden de produccion", style="Card.TLabelframe")
        top.pack(fill=tk.X)

        self.template_var = tk.StringVar()
        self.template_key_var = tk.StringVar()
        self.template_name_var = tk.StringVar()
        self.machine_type_var = tk.StringVar(value="cnc")
        self.consultation_count_var = tk.IntVar(value=5)
        self.line_var = tk.StringVar()
        self.program_var = tk.StringVar()
        self.product_code_var = tk.StringVar()
        self.product_name_var = tk.StringVar()
        self.line_mode_var = tk.StringVar(value="automatic")
        self.quality_checks_var = tk.StringVar()
        self.order_id_var = tk.StringVar()
        self.units_var = tk.StringVar()
        self.shift_var = tk.StringVar()
        self.protocol_var = tk.StringVar()
        self.operator_id_var = tk.StringVar()
        self.operator_name_var = tk.StringVar()
        self.status_var = tk.StringVar(value="Listo para ejecutar.")

        self._build_request_inputs(top)

        template_frame = ttk.LabelFrame(root, text="Configuracion de plantilla y producto", style="Card.TLabelframe")
        template_frame.pack(fill=tk.X, pady=(12, 0))
        self._build_template_inputs(template_frame)

        actions = ttk.Frame(top, style="Surface.TFrame")
        actions.grid(row=2, column=0, columnspan=10, sticky="ew", pady=(14, 0))
        actions.columnconfigure(9, weight=1)

        ttk.Button(actions, text="Nueva orden", command=self._reset_form).grid(row=0, column=0, padx=4)
        ttk.Button(actions, text="Guardar plantilla", command=self._save_template).grid(row=0, column=1, padx=4)
        ttk.Button(actions, text="Cargar plantilla", command=self._load_selected_template).grid(row=0, column=2, padx=4)
        ttk.Button(actions, text="Restablecer sistema", command=self._reset_context).grid(row=0, column=3, padx=4)
        ttk.Label(actions, text="Consultas", style="Field.TLabel").grid(row=0, column=4, padx=(8, 2))
        ttk.Spinbox(actions, from_=1, to=5, textvariable=self.consultation_count_var, width=5).grid(
            row=0, column=5, padx=(0, 4)
        )
        ttk.Button(actions, text="Generar consultas", style="Accent.TButton", command=self._generate_consultations).grid(
            row=0, column=6, padx=4
        )
        ttk.Button(actions, text="Exportar documento", command=self._export_consultations_document).grid(
            row=0, column=7, padx=8
        )
        ttk.Button(actions, text="Ejecutar demostracion", style="Primary.TButton", command=self._execute).grid(
            row=0, column=8, padx=8
        )
        ttk.Label(actions, textvariable=self.status_var, style="Status.TLabel").grid(
            row=0, column=9, sticky="e", padx=(12, 0)
        )

        body = ttk.Frame(root, style="App.TFrame")
        body.pack(fill=tk.BOTH, expand=True, pady=(12, 0))

        self.tabs = ttk.Notebook(body)
        self.tabs.pack(fill=tk.BOTH, expand=True)

        summary_tab = ttk.Frame(self.tabs, style="Surface.TFrame", padding=8)
        summary_tab.pack_propagate(False)
        self.summary_text = ScrolledText(summary_tab, wrap=tk.WORD, font=("Consolas", 10))
        self._style_text_panel(self.summary_text)
        self.summary_text.pack(fill=tk.BOTH, expand=True)
        self.tabs.add(summary_tab, text="Resumen general")

        results_tab = ttk.Frame(self.tabs, style="Surface.TFrame", padding=8)
        results_tab.pack_propagate(False)
        self.results_text = ScrolledText(results_tab, wrap=tk.WORD, font=("Consolas", 9))
        self._style_text_panel(self.results_text)
        self.results_text.pack(fill=tk.BOTH, expand=True)
        self.tabs.add(results_tab, text="Detalle de resultados")

        consultations_tab = ttk.Frame(self.tabs, style="Surface.TFrame", padding=8)
        consultations_tab.pack_propagate(False)
        self.consultations_text = ScrolledText(consultations_tab, wrap=tk.WORD, font=("Consolas", 9))
        self._style_text_panel(self.consultations_text)
        self.consultations_text.pack(fill=tk.BOTH, expand=True)
        self.tabs.add(consultations_tab, text="Consultas generadas")

    def _build_request_inputs(self, parent: ttk.LabelFrame) -> None:
        ttk.Label(parent, text="Plantilla guardada", style="Field.TLabel").grid(row=0, column=0, sticky="w", padx=4, pady=4)
        self.template_combo = ttk.Combobox(
            parent,
            textvariable=self.template_var,
            state="readonly",
            width=20,
        )
        self.template_combo.grid(row=1, column=0, sticky="ew", padx=4, pady=4)
        self.template_combo.bind("<<ComboboxSelected>>", self._on_template_selected)

        ttk.Label(parent, text="Codigo de orden", style="Field.TLabel").grid(row=0, column=1, sticky="w", padx=4, pady=4)
        self._create_placeholder_entry(parent, self.order_id_var, "Ej: OP-001", width=22).grid(
            row=1, column=1, sticky="ew", padx=4, pady=4
        )

        ttk.Label(parent, text="Unidades planeadas", style="Field.TLabel").grid(row=0, column=2, sticky="w", padx=4, pady=4)
        self._create_placeholder_entry(parent, self.units_var, "Ej: 300", width=14).grid(
            row=1, column=2, sticky="ew", padx=4, pady=4
        )

        ttk.Label(parent, text="Turno", style="Field.TLabel").grid(row=0, column=3, sticky="w", padx=4, pady=4)
        shift = ttk.Combobox(
            parent,
            textvariable=self.shift_var,
            values=["DAY", "NIGHT"],
            state="readonly",
            width=14,
        )
        shift.grid(row=1, column=3, sticky="ew", padx=4, pady=4)

        ttk.Label(parent, text="Protocolo", style="Field.TLabel").grid(row=0, column=4, sticky="w", padx=4, pady=4)
        protocol = ttk.Combobox(
            parent,
            textvariable=self.protocol_var,
            values=["opcua", "modbus"],
            state="readonly",
            width=14,
        )
        protocol.grid(row=1, column=4, sticky="ew", padx=4, pady=4)

        ttk.Label(parent, text="Codigo de usuario", style="Field.TLabel").grid(row=2, column=0, sticky="w", padx=4, pady=4)
        self._create_placeholder_entry(parent, self.operator_id_var, "Ej: USR-001", width=18).grid(
            row=3, column=0, sticky="ew", padx=4, pady=4
        )

        ttk.Label(parent, text="Nombre del usuario", style="Field.TLabel").grid(row=2, column=1, sticky="w", padx=4, pady=4)
        self._create_placeholder_entry(parent, self.operator_name_var, "Ej: Juan Sebastian", width=24).grid(
            row=3, column=1, columnspan=2, sticky="ew", padx=4, pady=4
        )

        for idx in range(5):
            parent.columnconfigure(idx, weight=1)

    def _build_template_inputs(self, parent: ttk.LabelFrame) -> None:
        ttk.Label(parent, text="Codigo de plantilla", style="Field.TLabel").grid(row=0, column=0, sticky="w", padx=4, pady=4)
        self._create_placeholder_entry(parent, self.template_key_var, "Ej: orden_demo", width=22).grid(
            row=1, column=0, sticky="ew", padx=4, pady=4
        )

        ttk.Label(parent, text="Nombre de plantilla", style="Field.TLabel").grid(row=0, column=1, sticky="w", padx=4, pady=4)
        self._create_placeholder_entry(parent, self.template_name_var, "Ej: ORDEN_DEMO", width=22).grid(
            row=1, column=1, sticky="ew", padx=4, pady=4
        )

        ttk.Label(parent, text="Tipo de maquina", style="Field.TLabel").grid(row=0, column=2, sticky="w", padx=4, pady=4)
        ttk.Combobox(
            parent,
            textvariable=self.machine_type_var,
            values=["cnc", "robot"],
            state="readonly",
            width=14,
        ).grid(row=1, column=2, sticky="ew", padx=4, pady=4)

        ttk.Label(parent, text="Linea", style="Field.TLabel").grid(row=0, column=3, sticky="w", padx=4, pady=4)
        self._create_placeholder_entry(parent, self.line_var, "Ej: line-a", width=18).grid(
            row=1, column=3, sticky="ew", padx=4, pady=4
        )

        ttk.Label(parent, text="Programa", style="Field.TLabel").grid(row=0, column=4, sticky="w", padx=4, pady=4)
        self._create_placeholder_entry(parent, self.program_var, "Ej: P-DEMO-01", width=18).grid(
            row=1, column=4, sticky="ew", padx=4, pady=4
        )

        ttk.Label(parent, text="Codigo de producto", style="Field.TLabel").grid(row=2, column=0, sticky="w", padx=4, pady=4)
        self._create_placeholder_entry(parent, self.product_code_var, "Ej: PRD-001", width=18).grid(
            row=3, column=0, sticky="ew", padx=4, pady=4
        )

        ttk.Label(parent, text="Nombre del producto", style="Field.TLabel").grid(row=2, column=1, sticky="w", padx=4, pady=4)
        self._create_placeholder_entry(parent, self.product_name_var, "Ej: Pieza Demo", width=24).grid(
            row=3, column=1, columnspan=2, sticky="ew", padx=4, pady=4
        )

        ttk.Label(parent, text="Modo de linea", style="Field.TLabel").grid(row=4, column=0, sticky="w", padx=4, pady=4)
        ttk.Combobox(
            parent,
            textvariable=self.line_mode_var,
            values=["automatic", "manual"],
            state="readonly",
            width=16,
        ).grid(row=5, column=0, sticky="ew", padx=4, pady=4)

        ttk.Label(parent, text="Controles de calidad (coma)", style="Field.TLabel").grid(row=4, column=1, columnspan=2, sticky="w", padx=4, pady=4)
        self._create_placeholder_entry(
            parent,
            self.quality_checks_var,
            "Ej: dimension_check,surface_check",
        ).grid(
            row=5, column=1, columnspan=2, sticky="ew", padx=4, pady=4
        )

        ttk.Label(parent, text="Parametros extra JSON", style="Field.TLabel").grid(row=4, column=3, columnspan=2, sticky="w", padx=4, pady=4)
        self.parameters_text = ScrolledText(parent, wrap=tk.WORD, font=("Consolas", 10), height=4)
        self.parameters_text.grid(row=5, column=3, columnspan=2, sticky="nsew", padx=4, pady=4)
        self._style_text_panel(self.parameters_text)
        self._bind_text_placeholder()

        hint = ttk.Label(
            parent,
            text="Producto y usuario son datos obligatorios del sistema, aunque no cuenten como evidencias separadas.",
            style="Hint.TLabel",
        )
        hint.grid(row=6, column=0, columnspan=5, sticky="w", padx=4, pady=(2, 0))

        for idx in range(5):
            parent.columnconfigure(idx, weight=1)

    def _style_text_panel(self, widget: ScrolledText) -> None:
        widget.configure(
            background=self.palette["text_panel"],
            foreground=self.palette["text"],
            insertbackground=self.palette["text"],
            selectbackground="#bfd4ea",
            selectforeground=self.palette["text"],
            borderwidth=1,
            relief="solid",
            highlightthickness=0,
            padx=14,
            pady=12,
            spacing1=1,
            spacing3=3,
            cursor="arrow",
            undo=False,
        )
        widget.bind("<Key>", lambda _event: "break")
        widget.bind("<<Paste>>", lambda _event: "break")
        widget.bind("<<Cut>>", lambda _event: "break")

    def _sync_template_options(self) -> None:
        template_keys = self.app.list_templates()
        self.template_combo.configure(values=template_keys)

    def _create_placeholder_entry(
        self,
        parent: ttk.LabelFrame,
        variable: tk.StringVar,
        placeholder: str,
        **kwargs,
    ) -> ttk.Entry:
        entry = ttk.Entry(parent, textvariable=variable, **kwargs)
        self._placeholder_vars[str(variable)] = placeholder
        self._placeholder_entries.append((entry, variable))
        entry.bind("<FocusIn>", lambda _event, widget=entry, var=variable: self._clear_placeholder(widget, var))
        entry.bind("<FocusOut>", lambda _event, widget=entry, var=variable: self._apply_placeholder(widget, var))
        self._apply_placeholder(entry, variable)
        return entry

    def _apply_placeholder(self, widget: ttk.Entry, variable: tk.StringVar) -> None:
        if variable.get().strip():
            return
        variable.set(self._placeholder_vars[str(variable)])
        widget.configure(foreground="#98a2b3")

    def _clear_placeholder(self, widget: ttk.Entry, variable: tk.StringVar) -> None:
        if variable.get() == self._placeholder_vars.get(str(variable), ""):
            variable.set("")
            widget.configure(foreground="#111827")

    def _is_placeholder_value(self, variable: tk.StringVar) -> bool:
        return variable.get() == self._placeholder_vars.get(str(variable), "")

    def _clean_value(self, variable: tk.StringVar) -> str:
        if self._is_placeholder_value(variable):
            return ""
        return variable.get()

    def _refresh_entry_styles(self) -> None:
        for entry, variable in self._placeholder_entries:
            if self._is_placeholder_value(variable):
                entry.configure(foreground="#98a2b3")
            else:
                entry.configure(foreground="#111827")

    def _apply_text_placeholder(self) -> None:
        raw_text = self.parameters_text.get("1.0", tk.END).strip()
        if raw_text:
            return
        self._text_placeholder_active = True
        self.parameters_text.configure(foreground="#98a2b3")
        self.parameters_text.delete("1.0", tk.END)
        self.parameters_text.insert("1.0", self._text_placeholder)

    def _clear_text_placeholder(self) -> None:
        if not self._text_placeholder_active:
            return
        self._text_placeholder_active = False
        self.parameters_text.configure(foreground="#111827")
        self.parameters_text.delete("1.0", tk.END)

    def _bind_text_placeholder(self) -> None:
        self.parameters_text.bind("<FocusIn>", lambda _event: self._clear_text_placeholder())
        self.parameters_text.bind("<FocusOut>", lambda _event: self._apply_text_placeholder())
        self._apply_text_placeholder()

    def _reset_form(self) -> None:
        self.template_var.set("")
        self.template_key_var.set("")
        self.template_name_var.set("")
        self.machine_type_var.set("cnc")
        self.line_var.set("")
        self.program_var.set("")
        self.product_code_var.set("")
        self.product_name_var.set("")
        self.line_mode_var.set("automatic")
        self.quality_checks_var.set("")
        self.order_id_var.set("")
        self.units_var.set("")
        self.shift_var.set("DAY")
        self.protocol_var.set("opcua")
        self.operator_id_var.set("")
        self.operator_name_var.set("")
        self.parameters_text.delete("1.0", tk.END)
        for entry, variable in self._placeholder_entries:
            variable.set("")
            self._apply_placeholder(entry, variable)
        self._apply_text_placeholder()
        self.last_request = None
        self.last_result = None
        self.saved_consultations = []
        self.consultation_batch_label = ""
        self._set_text(
            self.consultations_text,
            "Ejecuta la orden y luego pulsa 'Generar consultas' para registrarlas por separado.",
        )
        self.status_var.set("Formulario listo. Registra los datos y ejecuta la demostracion.")

    def _on_template_selected(self, _event: tk.Event | None = None) -> None:
        self._load_selected_template()

    def _load_selected_template(self) -> None:
        template_key = self.template_var.get().strip()
        if not template_key:
            messagebox.showwarning("MES UI", "Selecciona una plantilla guardada para cargarla.")
            return

        try:
            template = self.app.get_template_definition(template_key)
        except Exception as exc:
            messagebox.showerror("MES UI", str(exc))
            return

        self.template_key_var.set(template["template_key"])
        self.template_name_var.set(template["template_name"])
        self.machine_type_var.set(template["machine_type"])
        self.line_var.set(template["line"])
        self.program_var.set(template["program"])
        self.product_code_var.set(template["product_code"])
        self.product_name_var.set(template["product_name"])
        self.line_mode_var.set(template["line_mode"])
        self.quality_checks_var.set(", ".join(template["quality_checks"]))

        parameters = dict(template["parameters"])
        self._set_editor_text(self.parameters_text, json.dumps(parameters, ensure_ascii=False, indent=2))
        self._text_placeholder_active = False
        self.parameters_text.configure(foreground="#111827")
        self._refresh_entry_styles()
        self.saved_consultations = []
        self.consultation_batch_label = ""
        self._set_text(
            self.consultations_text,
            "Plantilla cargada. Ejecuta la orden y luego genera las consultas del proyecto.",
        )
        self.status_var.set(f"Plantilla '{template_key}' cargada correctamente.")

    def _reset_context(self) -> None:
        self.app = build_default_mes_project()
        self._sync_template_options()
        self._clear_outputs()
        self._reset_form()
        self.status_var.set("Sistema restablecido. Puedes registrar una nueva orden.")

    def _clear_outputs(self) -> None:
        self._set_text(self.summary_text, "")
        self._set_text(self.results_text, "")
        self._set_text(self.consultations_text, "")

    def _save_template(self) -> None:
        try:
            template_key = self._clean_value(self.template_key_var).strip()
            self.app.register_template(
                template_key=template_key,
                template_name=self._clean_value(self.template_name_var).strip(),
                machine_type=self.machine_type_var.get().strip(),
                line=self._clean_value(self.line_var).strip(),
                program=self._clean_value(self.program_var).strip(),
                product_code=self._clean_value(self.product_code_var).strip(),
                product_name=self._clean_value(self.product_name_var).strip(),
                quality_checks=self._parse_quality_checks(),
                line_mode=self.line_mode_var.get().strip(),
                parameters=self._parse_parameters(),
            )
            self._sync_template_options()
            self.template_var.set(template_key)
            self.saved_consultations = []
            self.consultation_batch_label = ""
            self._set_text(
                self.consultations_text,
                "Plantilla guardada. Ejecuta la orden y luego genera las consultas del proyecto.",
            )
            self.status_var.set(f"Plantilla '{template_key}' guardada correctamente.")
        except Exception as exc:
            self.status_var.set("No fue posible guardar la plantilla.")
            messagebox.showerror("MES UI", str(exc))

    def _execute(self) -> None:
        try:
            order_id = self._clean_value(self.order_id_var).strip()
            units = int(self._clean_value(self.units_var).strip())
            if units <= 0:
                raise ValueError("Planned units debe ser mayor que cero.")
            if not self.template_var.get().strip():
                raise ValueError("Selecciona o guarda una plantilla antes de ejecutar.")
            if not order_id:
                raise ValueError("Order ID es obligatorio.")
            if not order_id.startswith("OP-"):
                raise ValueError("Order ID debe empezar por 'OP-' para pasar la validación de inicio.")
            if not self._clean_value(self.operator_id_var).strip():
                raise ValueError("Operator ID es obligatorio.")
            if not self._clean_value(self.operator_name_var).strip():
                raise ValueError("Operator name es obligatorio.")

            request = OrderExecutionRequest(
                template_key=self.template_var.get().strip(),
                order_id=order_id,
                planned_units=units,
                shift=self.shift_var.get().strip(),
                protocol=self.protocol_var.get().strip(),
                operator_id=self._clean_value(self.operator_id_var).strip(),
                operator_name=self._clean_value(self.operator_name_var).strip(),
            )
            self.status_var.set("Ejecutando flujo de produccion...")
            self.update_idletasks()

            result = self.app.execute_order(request)
            self.last_request = request
            self.last_result = result
            self.saved_consultations = []
            self.consultation_batch_label = ""
            self._render_result(result)
            self.status_var.set("Demostracion completada correctamente.")
            self.tabs.select(0)
        except Exception as exc:
            self.status_var.set("Se presento un error durante la ejecucion.")
            messagebox.showerror("MES UI", str(exc))

    def _parse_quality_checks(self) -> list[str]:
        raw_checks = self._clean_value(self.quality_checks_var).strip()
        if not raw_checks:
            return []
        return [item.strip() for item in raw_checks.split(",") if item.strip()]

    def _parse_parameters(self) -> dict:
        if self._text_placeholder_active:
            return {}
        raw_content = self.parameters_text.get("1.0", tk.END).strip()
        if not raw_content:
            return {}

        try:
            parsed = json.loads(raw_content)
        except json.JSONDecodeError as exc:
            raise ValueError("Parametros extra JSON no tiene un formato válido.") from exc

        if not isinstance(parsed, dict):
            raise ValueError("Parametros extra JSON debe representar un objeto JSON.")
        return parsed

    def _render_result(self, result: dict) -> None:
        self._clear_outputs()

        summary_lines = []
        summary_lines.append(f"Status: {result.get('status')}")
        summary_lines.append(f"Order: {result.get('order_id', '-')}")
        if result.get("status") == "OK":
            operator = result.get("operator", {})
            product = result.get("product", {})
            summary_lines.append(
                f"Operator: {operator.get('operator_id', '-')} / {operator.get('operator_name', '-')}"
            )
            summary_lines.append(
                f"Product: {product.get('product_code', '-')} / {product.get('product_name', '-')}"
            )
            summary_lines.append(f"Dispatch: {result.get('dispatch')}")
            summary_lines.append(f"Legacy sync: {result.get('legacy_sync')}")
            summary_lines.append(f"Plant capacity units: {result.get('plant_capacity_units')}")
            secure_oee = result.get("secure_oee_report", {})
            telemetry = result.get("telemetry_summary", {})
            summary_lines.append(f"Secure OEE Proxy: {secure_oee.get('oee', '-')}%")
            summary_lines.append(
                "Flyweight telemetry: "
                f"events={telemetry.get('total_events', '-')}, "
                f"shared_profiles={telemetry.get('shared_profiles', '-')}, "
                f"saved={telemetry.get('estimated_objects_saved', '-')}"
            )
            controller = result.get("controller_summary", {})
            summary_lines.append("")
            summary_lines.append("Controller summary:")
            for key, value in controller.items():
                summary_lines.append(f"  - {key}: {value}")
        else:
            summary_lines.append(f"Message: {result.get('message')}")
        self._set_text(self.summary_text, "\n".join(summary_lines))

        self._render_results_detail(result)

    def _render_results_detail(self, result: dict) -> None:
        lines: list[str] = []
        lines.append("Contexto de la ejecucion")
        lines.append("")
        operator = result.get("operator", {})
        product = result.get("product", {})
        lines.append(f"Order ID: {result.get('order_id', '-')}")
        lines.append(f"Operator: {operator.get('operator_id', '-')} / {operator.get('operator_name', '-')}")
        lines.append(f"Product: {product.get('product_code', '-')} / {product.get('product_name', '-')}")
        lines.append("")

        lines.append("Resultado operativo del proyecto")
        lines.append("")
        lines.append(json.dumps(result.get("final_report", {}), ensure_ascii=False, indent=2))
        lines.append("")

        lines.append("OEE y auditoria")
        lines.append("")
        lines.append(
            json.dumps(
                {
                    "secure_oee_report": result.get("secure_oee_report", {}),
                    "proxy_audit_log": result.get("proxy_audit_log", []),
                },
                ensure_ascii=False,
                indent=2,
            )
        )
        lines.append("")

        lines.append("Telemetría")
        lines.append("")
        lines.append(json.dumps(result.get("telemetry_summary", {}), ensure_ascii=False, indent=2))
        lines.append("")

        lines.append("Trazabilidad de la ejecucion")
        lines.append("")
        for idx, row in enumerate(result.get("execution_log", []), start=1):
            lines.append(
                f"{idx:02d}. {row.get('step', '')} | {row.get('pattern', '')} | "
                f"{row.get('module', '')} | {row.get('evidence', '')}"
            )
        lines.append("")

        lines.append("Patrones utilizados en el flujo")
        lines.append("")
        for item in result.get("implementation_map", []):
            lines.append(
                f"- {item.get('pattern', '')}: {item.get('usage_point', '')} "
                f"({item.get('module', '')})"
            )

        self._set_text(self.results_text, "\n".join(lines))

    def _set_text(self, widget: ScrolledText, content: str) -> None:
        widget.delete("1.0", tk.END)
        widget.insert(tk.END, content)
        widget.mark_set("insert", "1.0")
        widget.yview_moveto(0.0)
        widget.xview_moveto(0.0)

    def _set_editor_text(self, widget: ScrolledText, content: str) -> None:
        widget.delete("1.0", tk.END)
        widget.insert(tk.END, content)

    def _get_consultation_count(self) -> int:
        count = int(self.consultation_count_var.get())
        if count <= 0:
            raise ValueError("El numero de consultas debe ser mayor que cero.")
        return count

    def _build_consultations(self) -> list[dict]:
        template_key = self.template_var.get().strip()
        if not template_key:
            raise ValueError("Selecciona una plantilla para generar las consultas.")
        if not self.last_request or not self.last_result or self.last_result.get("status") != "OK":
            raise ValueError("Ejecuta una orden valida antes de generar las consultas.")

        order_id = self._clean_value(self.order_id_var).strip() or "OP-PENDIENTE"
        units_raw = self._clean_value(self.units_var).strip()
        planned_units = int(units_raw) if units_raw else 0
        shift = self.shift_var.get().strip() or "DAY"
        protocol = self.protocol_var.get().strip() or "opcua"
        if self.last_request.order_id != order_id:
            raise ValueError("Las consultas deben generarse con la misma orden que acabas de ejecutar.")

        result_context = self.last_result if self.last_request.template_key == template_key else None
        consultation_count = self._get_consultation_count()
        return self.app.generate_production_consultations(
            template_key=template_key,
            order_id=order_id,
            planned_units=planned_units,
            shift=shift,
            protocol=protocol,
            operator_id=self._clean_value(self.operator_id_var).strip(),
            operator_name=self._clean_value(self.operator_name_var).strip(),
            result=result_context,
            consultation_count=consultation_count,
        )

    def _generate_consultations(self) -> None:
        try:
            consultations = self._build_consultations()
        except Exception as exc:
            self.status_var.set("No fue posible generar las consultas.")
            messagebox.showerror("MES UI", str(exc))
            return

        generated_at = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        self.consultation_batch_label = f"Lote generado: {generated_at}"
        self.saved_consultations = []
        for index, consultation in enumerate(consultations, start=1):
            self.saved_consultations.append(
                {
                    "record_number": index,
                    "generated_at": generated_at,
                    **consultation,
                }
            )
        self._render_consultations()
        consultation_count = len(self.saved_consultations)
        self.status_var.set(f"Se registraron {consultation_count} consultas para la orden ejecutada.")
        self.tabs.select(2)

    def _render_consultations(self) -> None:
        if not self.saved_consultations:
            self._set_text(
                self.consultations_text,
                "Aun no hay consultas registradas. Ejecuta la orden y luego pulsa 'Generar consultas'.",
            )
            return

        consultation_count = len(self.saved_consultations)
        lines: list[str] = []
        lines.append(f"{consultation_count} consultas registradas del proyecto de control de produccion")
        lines.append(self.consultation_batch_label)
        lines.append(f"Total de consultas guardadas en el lote actual: {consultation_count}")
        lines.append("Usuario y producto aparecen como contexto obligatorio del proceso, no como consultas aparte.")
        lines.append("")
        for consultation in self.saved_consultations:
            lines.append(f"Registro {consultation['record_number']} - {consultation['identifier']}: {consultation['title']}")
            lines.append(f"Objetivo: {consultation['objective']}")
            lines.append(f"Pregunta: {consultation['business_question']}")
            lines.append("Filtros:")
            for item in consultation["filters"]:
                lines.append(f"  - {item}")
            lines.append("Consulta sugerida:")
            lines.append(consultation["suggested_query"])
            lines.append("Patrones utilizados:")
            for pattern in consultation["pattern_details"]:
                lines.append(f"  - {pattern['pattern']}: {pattern['usage_point']}")
            lines.append("Resultado esperado:")
            lines.append(json.dumps(consultation["sample_result"], ensure_ascii=False, indent=2))
            lines.append("")
        self._set_text(self.consultations_text, "\n".join(lines))
        self.tabs.tab(2, text=f"Consultas generadas ({consultation_count})")

    def _export_consultations_document(self) -> None:
        if not self.saved_consultations:
            messagebox.showerror("MES UI", "Primero debes ejecutar la orden y generar las consultas.")
            return

        default_name = f"consultas_mes_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        initial_dir = Path(__file__).resolve().parent
        target_path = filedialog.asksaveasfilename(
            title="Guardar documento de evidencias del proyecto",
            defaultextension=".docx",
            initialdir=str(initial_dir),
            initialfile=default_name,
            filetypes=[("Word Document", "*.docx")],
        )
        if not target_path:
            return

        document = Document()
        consultation_count = len(self.saved_consultations)
        document.add_heading("Sistema MES - Evidencias funcionales del proyecto", level=0)
        document.add_paragraph(
            f"Documento generado desde la interfaz para sustentar {consultation_count} consultas del sistema de control de produccion "
            "y los patrones de software que soportan cada resultado."
        )

        context_lines = [
            f"Template: {self.template_var.get().strip() or '-'}",
            f"Product code: {self._clean_value(self.product_code_var).strip() or '-'}",
            f"Product name: {self._clean_value(self.product_name_var).strip() or '-'}",
            f"Order ID: {self._clean_value(self.order_id_var).strip() or '-'}",
            f"Planned units: {self._clean_value(self.units_var).strip() or '-'}",
            f"Shift: {self.shift_var.get().strip() or '-'}",
            f"Protocol: {self.protocol_var.get().strip() or '-'}",
            f"Operator ID: {self._clean_value(self.operator_id_var).strip() or '-'}",
            f"Operator name: {self._clean_value(self.operator_name_var).strip() or '-'}",
        ]
        document.add_heading("Contexto registrado", level=1)
        for line in context_lines:
            document.add_paragraph(line, style="List Bullet")

        document.add_paragraph(self.consultation_batch_label or "Lote actual de consultas")

        for consultation in self.saved_consultations:
            document.add_heading(f"{consultation['identifier']} - {consultation['title']}", level=1)
            document.add_paragraph(f"Objetivo: {consultation['objective']}")
            document.add_paragraph(f"Pregunta de negocio: {consultation['business_question']}")
            document.add_paragraph("Filtros clave:", style=None)
            for item in consultation["filters"]:
                document.add_paragraph(item, style="List Bullet")

            document.add_paragraph("Consulta sugerida:")
            document.add_paragraph(consultation["suggested_query"])

            document.add_paragraph("Patrones aplicados:")
            for pattern in consultation["pattern_details"]:
                document.add_paragraph(
                    f"{pattern['pattern']} - {pattern['objective']} ({pattern['usage_point']})",
                    style="List Bullet",
                )

            document.add_paragraph("Resultado esperado:")
            document.add_paragraph(json.dumps(consultation["sample_result"], ensure_ascii=False, indent=2))

        document.save(target_path)
        self.status_var.set(f"Documento generado: {Path(target_path).name}")
        messagebox.showinfo("MES UI", f"Documento generado correctamente:\n{target_path}")


def main() -> None:
    ui = MESMiniUI()
    ui.mainloop()


if __name__ == "__main__":
    main()
